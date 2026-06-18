#!/usr/bin/env python3
"""Rebuild Bao_cao_MHNNL theo cấu trúc sách Foundation of LLM (Xiao & Zhu, 2025)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon/Bao_cao_MHNNL_UTC_Assistant.docx"

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# Style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
for attr in ['w:ascii','w:hAnsi','w:cs','w:eastAsia']:
    rFonts.set(qn(attr), 'Times New Roman')
rPr.insert(0, rFonts)
sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26'); rPr.append(sz)

# Helper functions
def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        for attr in ['w:ascii','w:hAnsi','w:cs','w:eastAsia']:
            rFonts.set(qn(attr), 'Times New Roman')
        rPr.insert(0, rFonts)
    return h

def P(text, bold=False, align=None, size=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.font.size = Pt(size)
    if align is not None: p.alignment = align
    return p

def TBL(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    doc.add_paragraph()
    return table

def BR(): doc.add_page_break()
def SPC(n=1):
    for _ in range(n): doc.add_paragraph()

# ═══════════════════════════ TRANG BÌA ═══════════════════════════
SPC(4)
P("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
P("KHOA CÔNG NGHỆ THÔNG TIN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
SPC(3)
P("BÁO CÁO MÔN HỌC", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
P("MÔ HÌNH NGÔN NGỮ LỚN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
SPC(2)
P("ĐỀ TÀI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
P("XÂY DỰNG HỆ THỐNG TRỢ LÝ ẢO HỖ TRỢ SINH VIÊN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
P("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
P("SỬ DỤNG MÔ HÌNH NGÔN NGỮ LỚN VÀ RAG", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
SPC(2)
for line in ["Giảng viên hướng dẫn : TS. ……………………",
             "Sinh viên thực hiện  : Trần Văn Trung",
             "Mã sinh viên         : ……………………",
             "Lớp                  : ……………………",
             "Khóa                 : ……………………"]:
    P(line, size=13)
SPC(4)
P("Hà Nội – 2026", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
BR()

# ═══════════════════════════ LỜI CẢM ƠN ═══════════════════════════
H("LỜI CẢM ƠN", 1)
P("Trong quá trình học tập và thực hiện báo cáo môn học \"Mô hình ngôn ngữ lớn\", em đã nhận được sự hướng dẫn, giúp đỡ tận tình từ các thầy cô giáo trong Khoa Công nghệ Thông tin – Trường Đại học Giao thông Vận tải. Em xin gửi lời cảm ơn chân thành đến các thầy cô đã trang bị cho em những kiến thức nền tảng về xử lý ngôn ngữ tự nhiên, học sâu và đặc biệt là về các mô hình ngôn ngữ lớn (LLM) và kỹ thuật truy vấn tạo sinh tăng cường (RAG), làm cơ sở để em có thể thực hiện báo cáo này.")
P("Em cũng xin cảm ơn gia đình và bạn bè đã luôn động viên, hỗ trợ em trong suốt thời gian học tập và nghiên cứu. Mặc dù đã cố gắng hết sức, nhưng do hạn chế về thời gian và kiến thức, báo cáo không tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý của thầy cô để báo cáo được hoàn thiện hơn.")
P("Em xin trân trọng cảm ơn!", bold=True)
BR()

# ═══════════════════════════ MỤC LỤC ═══════════════════════════
H("MỤC LỤC", 1)
toc = [
    ("Lời cảm ơn", "2"),
    ("Mục lục", "3"),
    ("Danh mục các từ viết tắt", "5"),
    ("Danh mục bảng biểu", "6"),
    ("Danh mục hình ảnh", "7"),
    ("Mở đầu", "9"),
    ("Chương 1. Pre-training – Tiền huấn luyện", "12"),
    ("1.1. Các mô hình Pre-training trong NLP", "12"),
    ("1.2. Các tác vụ Self-supervised Pre-training", "15"),
    ("1.3. Ví dụ điển hình: BERT", "18"),
    ("1.4. Ứng dụng các mô hình BERT", "21"),
    ("Chương 2. Generative Models – Mô hình sinh", "23"),
    ("2.1. Giới thiệu về Mô hình Ngôn ngữ Lớn (LLM)", "23"),
    ("2.2. Huấn luyện ở quy mô lớn (Training at Scale)", "27"),
    ("2.3. Mô hình hóa chuỗi dài (Long Sequence Modeling)", "30"),
    ("Chương 3. Prompting – Kỹ thuật tạo prompt", "32"),
    ("3.1. Thiết kế Prompt cơ bản", "32"),
    ("3.2. Các phương pháp Prompting nâng cao", "36"),
    ("3.3. Learning to Prompt", "40"),
    ("Chương 4. Alignment – Căn chỉnh mô hình", "42"),
    ("4.1. Tổng quan về Alignment", "42"),
    ("4.2. Instruction Alignment (Căn chỉnh theo chỉ dẫn)", "43"),
    ("4.3. Human Preference Alignment: RLHF", "45"),
    ("4.4. Các cải tiến trong Human Preference Alignment", "47"),
    ("Chương 5. Inference – Suy luận", "49"),
    ("5.1. Prefilling và Decoding", "49"),
    ("5.2. Kỹ thuật Suy luận Hiệu quả (Efficient Inference)", "51"),
    ("5.3. Inference-time Scaling", "53"),
    ("Chương 6. Xây dựng Hệ thống UTC Assistant", "55"),
    ("6.1. Bài toán và Yêu cầu Hệ thống", "55"),
    ("6.2. Kiến trúc Tổng thể", "57"),
    ("6.3. Pipeline RAG 3 tầng", "59"),
    ("6.4. Xử lý Dữ liệu và Chunking", "63"),
    ("6.5. Giao diện Người dùng", "65"),
    ("Chương 7. Kết quả và Đánh giá", "67"),
    ("7.1. Đánh giá Chất lượng Truy vấn", "67"),
    ("7.2. Đánh giá Hiệu năng Hệ thống", "69"),
    ("7.3. So sánh với các Phương pháp khác", "71"),
    ("Kết luận và Kiến nghị", "73"),
    ("Danh mục Tài liệu Tham khảo", "75"),
]
for item, page in toc:
    dots = '.' * max(2, 65 - len(item))
    P(f"{item} {dots} {page}", size=12)
BR()

print("Part 1 done: Cover + TOC. Continuing...")
