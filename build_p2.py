#!/usr/bin/env python3
"""Part 2: Add Mở đầu + all 7 chapters + Kết luận + TLTK to existing .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

INPUT = "/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon/Bao_cao_MHNNL_UTC_Assistant.docx"
doc = Document(INPUT)

def H(t, lv=1):
    h = doc.add_heading(t, level=lv)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
    return h
def P(t, b=False, s=13):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = 'Times New Roman'; r.bold = b; r.font.size = Pt(s); return p
def TBL(hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs)); t.style = 'Table Grid'
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
    doc.add_paragraph(); return t
def BR(): doc.add_page_break()

# ═══════════════════════════════════ MỞ ĐẦU ═══════════════════════════════════
H("MỞ ĐẦU",1)
H("i. Bối cảnh",2)
P("Sự phát triển vượt bậc của các mô hình ngôn ngữ lớn (Large Language Models - LLM) trong những năm gần đây đã tạo ra một cuộc cách mạng trong lĩnh vực trí tuệ nhân tạo nói chung và xử lý ngôn ngữ tự nhiên nói riêng. Các mô hình như GPT-4, Claude, Gemini, Llama và Qwen đã chứng minh khả năng vượt trội trong việc hiểu và sinh ngôn ngữ tự nhiên, mở ra nhiều ứng dụng thực tiễn trong giáo dục, y tế, tài chính và dịch vụ khách hàng.")
P("Trường Đại học Giao thông Vận tải (UTC) với hàng chục nghìn sinh viên đang phải đối mặt với thách thức lớn trong việc hỗ trợ sinh viên tra cứu thông tin về quy chế đào tạo, học phí, học bổng, ký túc xá và các thủ tục hành chính. Việc xây dựng một hệ thống trợ lý ảo thông minh, hoạt động 24/7, có khả năng trả lời chính xác các câu hỏi của sinh viên bằng tiếng Việt là một nhu cầu cấp thiết.")
P("Báo cáo này được thực hiện trong khuôn khổ môn học \"Mô hình Ngôn ngữ Lớn\", nhằm nghiên cứu một cách có hệ thống về nền tảng lý thuyết của LLM – từ pre-training, generative models, prompting, alignment đến inference – đồng thời áp dụng các kiến thức này vào việc xây dựng hệ thống UTC Assistant, một chatbot hỗ trợ sinh viên sử dụng kỹ thuật Retrieval-Augmented Generation (RAG).")

H("ii. Phạm vi và Mục tiêu",2)
P("Phạm vi của báo cáo bao gồm hai phần chính:")
P("Phần 1 – Nền tảng Lý thuyết (Chương 1-5): Trình bày có hệ thống về 5 trụ cột của LLM hiện đại theo cấu trúc của cuốn sách \"Foundations of Large Language Models\" (Xiao & Zhu, 2025), bao gồm: Pre-training, Generative Models, Prompting, Alignment và Inference.")
P("Phần 2 – Ứng dụng Thực tiễn (Chương 6-7): Mô tả chi tiết quá trình thiết kế, xây dựng và đánh giá hệ thống UTC Assistant – trợ lý ảo hỗ trợ sinh viên Trường Đại học Giao thông Vận tải, sử dụng pipeline RAG 3 tầng với các mô hình embedding BAAI/bge-m3 và LLM qwen35-opus.")
P("Mục tiêu cụ thể của báo cáo:")
P("(1) Nắm vững nền tảng lý thuyết về LLM: hiểu được cách LLM được huấn luyện (pre-training), cách tạo prompt hiệu quả (prompting), cách căn chỉnh mô hình với mong đợi của con người (alignment), và cách tối ưu hóa quá trình suy luận (inference).")
P("(2) Xây dựng thành công hệ thống UTC Assistant: thiết kế kiến trúc, triển khai pipeline RAG 3 tầng, xử lý dữ liệu tiếng Việt, và đánh giá chất lượng hệ thống.")
P("(3) Đạt được các chỉ số chất lượng cao: Hit Rate 100%, MRR 0.980 trên bộ test 50 câu hỏi tiếng Việt.")

H("iii. Phương pháp Nghiên cứu",2)
P("Báo cáo sử dụng kết hợp phương pháp nghiên cứu lý thuyết và thực nghiệm:")
P("• Nghiên cứu lý thuyết: Tổng hợp, phân tích các công trình nghiên cứu nền tảng về LLM từ các hội nghị và tạp chí uy tín (NeurIPS, ICML, ACL, EMNLP) và các tài liệu chuyên khảo (Xiao & Zhu, 2025; Vaswani et al., 2017; Devlin et al., 2019; Brown et al., 2020; Lewis et al., 2020).")
P("• Thiết kế hệ thống: Áp dụng kiến trúc microservices (FastAPI + Next.js), thiết kế API RESTful, sử dụng mô hình embedding đa ngôn ngữ BAAI/bge-m3 (Chen et al., 2024) và cơ sở dữ liệu vector ChromaDB.")
P("• Thực nghiệm: Xây dựng hệ thống thực tế, đánh giá trên tập dữ liệu thực (Sổ tay sinh viên K66, 92 trang), đo lường các chỉ số Hit Rate, MRR, thời gian phản hồi và tỉ lệ lỗi.")

H("iv. Bố cục Báo cáo",2)
P("Báo cáo được tổ chức thành 7 chương, chia làm 2 phần:")
P("Phần 1 – Nền tảng Lý thuyết (Chương 1-5):")
P("• Chương 1: Pre-training – Trình bày các phương pháp tiền huấn luyện mô hình ngôn ngữ, bao gồm các tác vụ self-supervised và case study về BERT.")
P("• Chương 2: Generative Models – Giới thiệu về LLM, kiến trúc Decoder-only Transformer, quy trình huấn luyện quy mô lớn và các kỹ thuật mô hình hóa chuỗi dài.")
P("• Chương 3: Prompting – Trình bày các kỹ thuật thiết kế prompt từ cơ bản đến nâng cao, bao gồm In-Context Learning, Chain-of-Thought, và RAG.")
P("• Chương 4: Alignment – Thảo luận về các phương pháp căn chỉnh LLM với giá trị và mong đợi của con người, bao gồm RLHF và DPO.")
P("• Chương 5: Inference – Phân tích quy trình suy luận của LLM, các kỹ thuật tối ưu hóa và inference-time scaling.")
P("Phần 2 – Ứng dụng Thực tiễn (Chương 6-7):")
P("• Chương 6: Xây dựng Hệ thống UTC Assistant – Mô tả chi tiết kiến trúc, pipeline RAG 3 tầng, xử lý dữ liệu và giao diện người dùng.")
P("• Chương 7: Kết quả và Đánh giá – Phân tích kết quả thực nghiệm, đánh giá chất lượng retrieval và hiệu năng hệ thống.")
BR()

print("Done: Mở đầu. Saving intermediate...")
doc.save(INPUT)
print(f"Size: {os.path.getsize(INPUT):,} bytes")
