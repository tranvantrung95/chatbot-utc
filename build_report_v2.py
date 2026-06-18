#!/usr/bin/env python3
"""Generate bao cao MHNNL - Cau truc UTC tu van noi quy (v2)
Theo dung cau truc nguoi dung yeu cau: 4 chuong + Mo dau + Ket luan + TLTK
Font: Times New Roman 13pt, line spacing 1.5
"""
import os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon/Bao_cao_MHNNL_UTC_Assistant.docx"
doc = Document()

for sec in doc.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2.5)

# Style setup: Times New Roman 13pt, line spacing 1.5, justify, first-line indent 1.27cm
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(13)
sty.paragraph_format.line_spacing = 1.5
sty.paragraph_format.space_after = Pt(6)
sty.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
sty.paragraph_format.first_line_indent = Cm(1.27)
rPr = sty.element.get_or_add_rPr()
rf = OxmlElement('w:rFonts')
for a in ['w:ascii','w:hAnsi','w:cs','w:eastAsia']: rf.set(qn(a), 'Times New Roman')
rPr.insert(0, rf)
sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26'); rPr.append(sz)

for lv in [1,2,3]:
    hs = doc.styles[f'Heading {lv}']
    hs.font.name = 'Times New Roman'; hs.font.color.rgb = RGBColor(0,0,0)
    hs.font.bold = True
    sizes = {1:16, 2:14, 3:13}
    hs.font.size = Pt(sizes[lv])

def H(t, lv=1):
    h = doc.add_heading(t, level=lv)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
    return h

def P(t, b=False, s=13, align=None, indent=True):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = 'Times New Roman'; r.bold = b; r.font.size = Pt(s)
    # Explicit formatting: justify + 1.5 spacing + first-line indent
    pf = p.paragraph_format
    pf.alignment = align if align else WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Cm(1.27)
    return p

def PC(t, s=13):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = 'Times New Roman'; r.font.size = Pt(s)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; return p

def TBL(hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs)); t.style = 'Table Grid'
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
    doc.add_paragraph(); return t

DIR = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS = os.path.join(DIR, 'diagrams')

def BR(): doc.add_page_break()
def SP(n=1):
    for _ in range(n): doc.add_paragraph()
def IMG(filename, caption, width=5.5):
    """Embed a PNG image with centered caption"""
    path = os.path.join(DIAGRAMS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Inches(width))
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.font.name = 'Times New Roman'
        cr.font.size = Pt(11)
        cr.italic = True
        doc.add_paragraph()
    else:
        print(f'WARNING: Image not found: {path}')

print("=== GENERATING BAO CAO MHNNL - UTC TU VAN NOI QUY ===")

# ═══════════════════════════ COVER ═══════════════════════════
SP(4)
PC("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", s=14)
PC("KHOA CÔNG NGHỆ THÔNG TIN", s=14)
SP(3)
PC("BÁO CÁO MÔN HỌC", s=16)
PC("MÔ HÌNH NGÔN NGỮ LỚN", s=14)
SP(2)
PC("ĐỀ TÀI", s=14)
PC("XÂY DỰNG HỆ THỐNG TƯ VẤN NỘI QUY", s=13)
PC("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", s=13)
PC("SỬ DỤNG MÔ HÌNH NGÔN NGỮ LỚN VÀ RAG", s=13)
SP(2)
for ln in ["Giảng viên hướng dẫn : TS. ……………………",
           "Sinh viên thực hiện  : Trần Văn Trung",
           "Mã sinh viên         : ……………………",
           "Lớp                  : ……………………",
           "Khóa                 : ……………………"]:
    P(ln, s=13)
SP(4)
PC("Hà Nội – 2026", s=14)
BR()

# ═══════════════════════════ LỜI CẢM ƠN ═══════════════════════════
H("LỜI CẢM ƠN", 1)
P("Trong quá trình học tập và thực hiện báo cáo môn học \"Mô hình ngôn ngữ lớn\", em đã nhận được sự hướng dẫn, giúp đỡ tận tình từ các thầy cô giáo trong Khoa Công nghệ Thông tin – Trường Đại học Giao thông Vận tải. Em xin gửi lời cảm ơn chân thành đến các thầy cô đã trang bị cho em những kiến thức nền tảng về xử lý ngôn ngữ tự nhiên, học sâu và đặc biệt là về các mô hình ngôn ngữ lớn (LLM) và kỹ thuật truy vấn tạo sinh tăng cường (RAG).")
P("Em cũng xin cảm ơn gia đình và bạn bè đã luôn động viên, hỗ trợ em trong suốt thời gian học tập và nghiên cứu. Mặc dù đã cố gắng hết sức, báo cáo không tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý của thầy cô để báo cáo được hoàn thiện hơn.")
P("Em xin trân trọng cảm ơn!")
BR()

# ═══════════════════════════ MỤC LỤC ═══════════════════════════
H("MỤC LỤC", 1)
toc_items = [
    ("Lời cảm ơn","3"),("Mục lục","4"),("Danh mục các từ viết tắt","6"),
    ("Danh mục bảng biểu","7"),("Danh mục hình ảnh","8"),("Mở đầu","10"),
    ("Chương 1. Tổng quan bài toán xây dựng hệ thống tư vấn nội quy trường Đại học Giao thông Vận tải","12"),
    ("  1.1. Đặt vấn đề","12"),
    ("  1.2. Mục tiêu và phạm vi đồ án","13"),
    ("  1.3. Công nghệ và công cụ sử dụng","13"),
    ("    1.3.1. Công nghệ sử dụng","13"),
    ("    1.3.2. Công cụ sử dụng","14"),
    ("Chương 2. Mô hình ngôn ngữ lớn","15"),
    ("  2.1. Tổng quan mô hình ngôn ngữ lớn","15"),
    ("  2.2. Kiến trúc Transformer","16"),
    ("    2.2.1. Tổng quan kiến trúc","16"),
    ("    2.2.2. Encoder – Decoder","17"),
    ("    2.2.3. Cơ chế self-attention","18"),
    ("  2.3. Prompt Engineering","19"),
    ("Chương 3. Truy vấn tạo sinh tăng cường","21"),
    ("  3.1. Tổng quan về RAG","21"),
    ("  3.2. Phân loại RAG","22"),
    ("  3.3. Quy trình thực hiện phương pháp RAG ngây thơ (Naive RAG)","23"),
    ("    3.3.1. Indexing","24"),
    ("    3.3.2. Nhúng từ (Embedding)","26"),
    ("    3.3.3. Retrieving","27"),
    ("    3.3.4. Generating","28"),
    ("Chương 4. Xây dựng hệ thống tư vấn nội quy trường Đại học Giao thông Vận tải","30"),
    ("  4.1. Phân tích thiết kế hệ thống","30"),
    ("    4.1.1. Sơ đồ ca sử dụng","30"),
    ("    4.1.2. Đặc tả các chức năng","31"),
    ("  4.2. Xây dựng trang web tư vấn viên nội quy","35"),
    ("    4.2.1. Cài đặt chatbot sử dụng phương pháp RAG","35"),
    ("    4.2.2. Phân lớp câu hỏi của người dùng","39"),
    ("    4.2.3. Giao diện trang web tư vấn viên nội quy","40"),
    ("  4.3. Xây dựng trang web quản trị tư vấn nội quy","41"),
    ("    4.3.1. Giao diện chức năng xác thực","41"),
    ("    4.3.2. Giao diện chức năng thống kê","43"),
    ("    4.3.3. Giao diện chức năng cập nhật dữ liệu","44"),
    ("KẾT LUẬN VÀ KIẾN NGHỊ","47"),
    ("DANH MỤC TÀI LIỆU THAM KHẢO","49"),
]
for item, page in toc_items:
    P(f"{item} {'.' * max(2, 72 - len(item))} {page}", s=12, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
BR()

# ═══════════════════════════ DANH MỤC TỪ VIẾT TẮT ═══════════════════════════
H("DANH MỤC CÁC TỪ VIẾT TẮT", 1)
TBL(["Dạng viết tắt", "Dạng đầy đủ", "Ý nghĩa"], [
    ["LLM","Large Language Model","Mô hình ngôn ngữ lớn"],
    ["RAG","Retrieval-Augmented Generation","Truy vấn tạo sinh tăng cường"],
    ["NLP","Natural Language Processing","Xử lý ngôn ngữ tự nhiên"],
    ["BERT","Bidirectional Encoder Representations from Transformers","Biểu diễn mã hóa hai chiều từ Transformer"],
    ["GPT","Generative Pre-trained Transformer","Transformer sinh tiền huấn luyện"],
    ["CoT","Chain of Thought","Chuỗi suy luận"],
    ["ICL","In-Context Learning","Học trong ngữ cảnh"],
    ["MLM","Masked Language Modeling","Mô hình hóa ngôn ngữ dạng mặt nạ"],
    ["MRR","Mean Reciprocal Rank","Xếp hạng đảo ngược trung bình"],
    ["BM25","Best Matching 25","Thuật toán xếp hạng văn bản"],
    ["RRF","Reciprocal Rank Fusion","Kết hợp xếp hạng đảo ngược"],
    ["SSE","Server-Sent Events","Sự kiện gửi từ máy chủ"],
    ["OCR","Optical Character Recognition","Nhận dạng ký tự quang học"],
    ["API","Application Programming Interface","Giao diện lập trình ứng dụng"],
    ["UTC","University of Transport and Communications","Trường Đại học Giao thông Vận tải"],
    ["DB","Database","Cơ sở dữ liệu"],
    ["FE","Frontend","Giao diện người dùng"],
    ["BE","Backend","Máy chủ xử lý"],
    ["RoPE","Rotary Position Embedding","Nhúng vị trí xoay"],
    ["GQA","Grouped Query Attention","Tự chú ý nhóm truy vấn"],
    ["NER","Named Entity Recognition","Nhận dạng thực thể có tên"],
    ["QA","Question Answering","Hỏi đáp"],
])
BR()

# ═══════════════════════════ DANH MỤC BẢNG BIỂU ═══════════════════════════
H("DANH MỤC BẢNG BIỂU", 1)
for t in [
    "Bảng 1.1: Các công nghệ chính sử dụng trong hệ thống",
    "Bảng 1.2: Các công cụ phát triển và triển khai",
    "Bảng 2.1: Một số mô hình LLM tiêu biểu",
    "Bảng 2.2: So sánh các chiến lược Prompt Engineering",
    "Bảng 3.1: So sánh các phân loại RAG",
    "Bảng 3.2: Các phương pháp embedding tiêu biểu",
    "Bảng 3.3: So sánh các chiến lược retrieval",
    "Bảng 4.1: Đặc tả ca sử dụng – Tra cứu nội quy",
    "Bảng 4.2: Đặc tả ca sử dụng – Quản lý dữ liệu",
    "Bảng 4.3: Đặc tả ca sử dụng – Thống kê và báo cáo",
    "Bảng 4.4: Cấu trúc metadata trong mỗi chunk của ChromaDB",
    "Bảng 4.5: Các API endpoint chính của UTC Assistant",
    "Bảng 4.6: Các lớp câu hỏi và chiến lược xử lý",
    "Bảng 4.7: Các quy tắc nghiệp vụ chính trong hệ thống",
]:
    P(f"{t} {'·' * max(2, 68 - len(t))} Trang", s=12, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
BR()

# ═══════════════════════════ DANH MỤC HÌNH ẢNH ═══════════════════════════
H("DANH MỤC HÌNH ẢNH", 1)
for f in [
    "Hình 1.1: Sơ đồ tổng quan bài toán hệ thống tư vấn nội quy UTC",
    "Hình 2.1: Kiến trúc Transformer (Vaswani et al., 2017)",
    "Hình 2.2: Cơ chế Scaled Dot-Product Attention",
    "Hình 2.3: Cơ chế Multi-Head Attention",
    "Hình 2.4: Sơ đồ minh họa các chiến lược Prompt Engineering",
    "Hình 3.1: Kiến trúc tổng quan RAG (Lewis et al., 2020)",
    "Hình 3.2: Phân loại các biến thể RAG",
    "Hình 3.3: Quy trình Naive RAG: Indexing → Retrieval → Generation",
    "Hình 3.4: Minh họa quá trình embedding văn bản thành vector",
    "Hình 4.1: Sơ đồ ca sử dụng tổng quan hệ thống UTC Assistant",
    "Hình 4.2: Sơ đồ ca sử dụng chi tiết – Nhóm chức năng sinh viên",
    "Hình 4.3: Sơ đồ ca sử dụng chi tiết – Nhóm chức năng quản trị",
    "Hình 4.4: Kiến trúc tổng thể hệ thống UTC Assistant",
    "Hình 4.5: Pipeline RAG 3 tầng: Bi-encoder → BM25 → LLM Reranker",
    "Hình 4.6: Luồng xử lý dữ liệu: PDF → OCR → Chunk → ChromaDB",
    "Hình 4.7: Giao diện chatbot tư vấn nội quy (trang sinh viên)",
    "Hình 4.8: Giao diện đăng nhập trang quản trị",
    "Hình 4.9: Giao diện Dashboard thống kê – Trang quản trị",
    "Hình 4.10: Giao diện cập nhật dữ liệu – Trang quản lý tài liệu",
]:
    P(f"{f} {'·' * max(2, 68 - len(f))} Trang", s=12, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
BR()

# ═══════════════════════════ MỞ ĐẦU ═══════════════════════════
H("MỞ ĐẦU", 1)

H("i. Bối cảnh và lý do chọn đề tài", 2)
P("Trong bối cảnh chuyển đổi số diễn ra mạnh mẽ trên mọi lĩnh vực, giáo dục đại học cũng không nằm ngoài xu hướng này. Trường Đại học Giao thông Vận tải (UTC) với quy mô hàng chục nghìn sinh viên đang phải đối mặt với thách thức lớn trong việc hỗ trợ sinh viên tra cứu thông tin về nội quy, quy chế của nhà trường. Sổ tay sinh viên dày 92 trang chứa đựng lượng lớn quy chế, quy định quan trọng về đào tạo, học phí, kỷ luật, ký túc xá, bảo hiểm y tế... nhưng sinh viên thường gặp khó khăn khi tra cứu do tài liệu dài, cấu trúc phức tạp và không có công cụ tìm kiếm ngữ nghĩa hiệu quả.")
P("Sự phát triển vượt bậc của các mô hình ngôn ngữ lớn (LLM) trong những năm gần đây đã mở ra những khả năng mới trong việc xây dựng các hệ thống hỏi đáp thông minh. Bắt đầu từ kiến trúc Transformer được giới thiệu bởi Vaswani et al. (2017) [1], tiếp đó là sự ra đời của BERT (Devlin et al., 2019) [2] và dòng GPT (Brown et al., 2020) [3], các LLM đã chứng minh khả năng vượt trội trong việc hiểu và sinh ngôn ngữ tự nhiên. Kết hợp LLM với kỹ thuật Truy vấn tạo sinh tăng cường (Retrieval-Augmented Generation - RAG) do Lewis et al. (2020) [4] đề xuất cho phép xây dựng các hệ thống có khả năng trả lời câu hỏi dựa trên tài liệu tham khảo cụ thể, giảm thiểu hiện tượng \"ảo giác\" (hallucination) và đảm bảo câu trả lời có căn cứ.")
P("Xuất phát từ nhu cầu thực tế đó, đề tài \"Xây dựng hệ thống tư vấn nội quy trường Đại học Giao thông Vận tải sử dụng mô hình ngôn ngữ lớn và RAG\" được lựa chọn thực hiện. Hệ thống hướng tới mục tiêu trở thành một trợ lý ảo thông minh, hoạt động 24/7, có khả năng trả lời chính xác các câu hỏi bằng tiếng Việt về nội quy, quy chế nhà trường.")

H("ii. Mục tiêu, đối tượng và phạm vi nghiên cứu", 2)
P("Mục tiêu tổng quát của đồ án là xây dựng một hệ thống trợ lý ảo (chatbot) có khả năng tư vấn, giải đáp các câu hỏi của sinh viên liên quan đến nội quy, quy chế của Trường Đại học Giao thông Vận tải một cách chính xác, nhanh chóng. Các mục tiêu cụ thể bao gồm:")
P("(1) Nghiên cứu và nắm vững nền tảng lý thuyết về mô hình ngôn ngữ lớn, bao gồm kiến trúc Transformer, cơ chế self-attention, và các kỹ thuật Prompt Engineering.")
P("(2) Nghiên cứu và ứng dụng kỹ thuật Truy vấn tạo sinh tăng cường (RAG) để xây dựng pipeline xử lý tài liệu và truy xuất thông tin từ cơ sở dữ liệu vector.")
P("(3) Thiết kế và xây dựng hệ thống hoàn chỉnh bao gồm: chatbot tư vấn nội quy cho sinh viên (giao diện web) và trang web quản trị cho cán bộ quản lý (xác thực, thống kê, cập nhật dữ liệu).")
P("(4) Đánh giá chất lượng và hiệu năng hệ thống trên dữ liệu thực tế của nhà trường, bao gồm độ chính xác của câu trả lời, tốc độ phản hồi và trải nghiệm người dùng.")
P("Đối tượng nghiên cứu của đồ án bao gồm: (1) Mô hình ngôn ngữ lớn (LLM) và kiến trúc Transformer; (2) Kỹ thuật Truy vấn tạo sinh tăng cường (RAG); (3) Các phương pháp embedding văn bản và truy xuất thông tin ngữ nghĩa; (4) Công nghệ xây dựng ứng dụng web (FastAPI, Next.js).")
P("Phạm vi nghiên cứu: Đồ án tập trung vào việc xây dựng hệ thống tư vấn dựa trên dữ liệu Sổ tay sinh viên hệ đại học chính quy (K66, năm học 2024-2025) của Trường Đại học Giao thông Vận tải. Hệ thống hỗ trợ tiếng Việt, bao gồm hai giao diện chính: chatbot cho sinh viên (trang student) và trang quản trị cho cán bộ (trang admin). Dữ liệu đầu vào là file PDF Sổ tay sinh viên 92 trang, được xử lý qua pipeline OCR và tổ chức thành các chunk trong cơ sở dữ liệu vector ChromaDB.")

H("iii. Phương pháp nghiên cứu", 2)
P("Báo cáo sử dụng kết hợp phương pháp nghiên cứu lý thuyết và thực nghiệm. Về lý thuyết, báo cáo tổng hợp và phân tích các công trình nghiên cứu nền tảng từ các hội nghị hàng đầu (NeurIPS, ICML, ACL) và tài liệu chuyên khảo, đặc biệt là cuốn \"Foundations of Large Language Models\" của Xiao và Zhu (2025) [5]. Về thực nghiệm, báo cáo xây dựng hệ thống thực tế trên nền tảng FastAPI (Python 3.14) và Next.js 15, sử dụng mô hình embedding BAAI/bge-m3 (Chen et al., 2024) [6] và LLM qwen35-opus, đánh giá trên tập dữ liệu Sổ tay sinh viên K66 (92 trang) với 50 câu hỏi kiểm thử.")

H("iv. Bố cục báo cáo", 2)
P("Ngoài phần Mở đầu, Kết luận và Danh mục tài liệu tham khảo, báo cáo được tổ chức thành 4 chương. Chương 1 trình bày tổng quan về bài toán xây dựng hệ thống tư vấn nội quy, bao gồm đặt vấn đề, mục tiêu, phạm vi và các công nghệ sử dụng. Chương 2 đi sâu vào nền tảng lý thuyết về mô hình ngôn ngữ lớn, kiến trúc Transformer và kỹ thuật Prompt Engineering. Chương 3 trình bày về kỹ thuật Truy vấn tạo sinh tăng cường (RAG) và quy trình thực hiện Naive RAG. Chương 4 mô tả chi tiết quá trình phân tích, thiết kế và xây dựng hệ thống tư vấn nội quy UTC, bao gồm cả giao diện sinh viên và giao diện quản trị.")
BR()

print("Cover + Front matter + Mo dau done.")

# ╔══════════════════════════════════════════════════════════════╗
# ║              CHƯƠNG 1: TỔNG QUAN BÀI TOÁN                   ║
# ╚══════════════════════════════════════════════════════════════╝

H("CHƯƠNG 1. TỔNG QUAN BÀI TOÁN XÂY DỰNG HỆ THỐNG TƯ VẤN NỘI QUY TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", 1)

H("1.1. Đặt vấn đề", 2)
P("Trường Đại học Giao thông Vận tải (UTC) là một trong những trường đại học kỹ thuật hàng đầu tại Việt Nam, với lịch sử hơn 60 năm đào tạo. Hàng năm, nhà trường tiếp nhận hàng nghìn tân sinh viên, mỗi sinh viên đều được cung cấp Sổ tay sinh viên – tài liệu tổng hợp toàn bộ các quy chế, quy định liên quan đến quá trình học tập và rèn luyện tại trường. Tuy nhiên, thực tế cho thấy việc tra cứu thông tin từ Sổ tay sinh viên gặp nhiều khó khăn.")
P("Thứ nhất, Sổ tay sinh viên là một văn bản dài (92 trang đối với khóa K66), bao gồm nhiều chương mục với các quy định thuộc nhiều lĩnh vực khác nhau như đào tạo, tài chính, kỷ luật, ký túc xá, bảo hiểm y tế... Sinh viên thường không biết chính xác thông tin mình cần nằm ở phần nào, dẫn đến mất nhiều thời gian tìm kiếm.")
P("Thứ hai, nhiều quy định được viết bằng ngôn ngữ hành chính, pháp lý, gây khó hiểu cho sinh viên, đặc biệt là tân sinh viên năm nhất. Sinh viên cần một công cụ không chỉ tìm kiếm thông tin mà còn có khả năng diễn giải, giải thích các quy định một cách dễ hiểu.")
P("Thứ ba, các kênh hỗ trợ truyền thống như Phòng Công tác Chính trị - Sinh viên, cố vấn học tập hay giáo viên chủ nhiệm hoạt động theo giờ hành chính, không thể đáp ứng nhu cầu tra cứu 24/7 của sinh viên trong thời đại số.")
P("Sự xuất hiện của các mô hình ngôn ngữ lớn (LLM) với khả năng hiểu và sinh ngôn ngữ tự nhiên đã mở ra một hướng giải quyết mới cho bài toán này. Tuy nhiên, LLM thuần túy có một hạn chế lớn: chúng được huấn luyện trên dữ liệu Internet tổng quát và không có kiến thức chuyên sâu về các quy định cụ thể của UTC. Hơn nữa, LLM có xu hướng \"ảo giác\" (hallucination) – tạo ra thông tin nghe có vẻ hợp lý nhưng không chính xác. Đây là rủi ro không thể chấp nhận được trong bối cảnh tư vấn nội quy, nơi mà thông tin sai lệch có thể gây hậu quả nghiêm trọng cho sinh viên.")
P("Kỹ thuật Truy vấn tạo sinh tăng cường (RAG) giải quyết hạn chế này bằng cách kết hợp LLM với một cơ sở dữ liệu bên ngoài: trước khi sinh câu trả lời, hệ thống truy xuất các đoạn văn bản liên quan từ tài liệu tham khảo và cung cấp chúng như ngữ cảnh cho LLM. Nhờ đó, câu trả lời không chỉ dựa trên kiến thức tổng quát của LLM mà còn được neo (anchor) vào các tài liệu cụ thể, đảm bảo tính chính xác và có căn cứ.")

H("1.2. Mục tiêu và phạm vi đồ án", 2)
P("Mục tiêu tổng quát của đồ án là xây dựng thành công hệ thống UTC Assistant – một trợ lý ảo thông minh có khả năng tư vấn, giải đáp các câu hỏi liên quan đến nội quy, quy chế của Trường Đại học Giao thông Vận tải. Hệ thống hoạt động 24/7, hỗ trợ tiếng Việt, cung cấp câu trả lời chính xác dựa trên dữ liệu Sổ tay sinh viên của nhà trường.")
P("Các mục tiêu cụ thể:")
P("(1) Xây dựng pipeline xử lý dữ liệu hoàn chỉnh: từ file PDF Sổ tay sinh viên, qua OCR, trích xuất cấu trúc, chia nhỏ thành các chunk và lưu trữ vào cơ sở dữ liệu vector ChromaDB với metadata phong phú.")
P("(2) Xây dựng pipeline RAG 3 tầng: Bi-encoder retrieval → BM25 hybrid → LLM Reranker, đảm bảo truy xuất chính xác các chunk liên quan nhất đến câu hỏi của người dùng.")
P("(3) Xây dựng giao diện chatbot thân thiện cho sinh viên, hỗ trợ streaming phản hồi theo thời gian thực (SSE), hiển thị nguồn tham khảo và gợi ý câu hỏi.")
P("(4) Xây dựng giao diện quản trị với các chức năng: xác thực người dùng, thống kê lượt hỏi và chất lượng phản hồi, cập nhật tài liệu và quản lý cơ sở tri thức.")
P("Phạm vi đồ án: Hệ thống được xây dựng dựa trên dữ liệu Sổ tay sinh viên hệ đại học chính quy khóa K66 (năm học 2024-2025). Hệ thống bao gồm hai phân hệ chính: (1) Trang web tư vấn viên nội quy dành cho sinh viên, cung cấp giao diện chatbot để đặt câu hỏi và nhận câu trả lời; (2) Trang web quản trị dành cho cán bộ, cung cấp các chức năng quản lý và giám sát hệ thống.")

H("1.3. Công nghệ và công cụ sử dụng", 2)

H("1.3.1. Công nghệ sử dụng", 3)
P("Hệ thống UTC Assistant được xây dựng dựa trên các công nghệ hiện đại, được lựa chọn dựa trên các tiêu chí: hiệu năng, khả năng mở rộng, hỗ trợ tiếng Việt và tính cộng đồng (mã nguồn mở).")
P("Bảng 1.1: Các công nghệ chính sử dụng trong hệ thống", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Thành phần", "Công nghệ", "Phiên bản", "Vai trò"], [
    ["Backend Framework","FastAPI (Python)","3.14","Xử lý API, pipeline RAG, streaming SSE"],
    ["Frontend Framework","Next.js (React)","15","Giao diện người dùng, SSR, routing"],
    ["Vector Database","ChromaDB","0.5+","Lưu trữ và truy xuất vector embedding"],
    ["Embedding Model","BAAI/bge-m3","—","Tạo vector embedding đa ngôn ngữ (1024 chiều)"],
    ["LLM","Qwen 3.5 (Opus)","—","Sinh câu trả lời tiếng Việt, context 128K token"],
    ["OCR Engine","Marker-PDF","—","Trích xuất văn bản từ PDF, giữ cấu trúc"],
    ["Database","SQLite","3","Lưu trữ câu hỏi, hoạt động, cấu hình hệ thống"],
    ["CSS Framework","Tailwind CSS","4","Thiết kế giao diện responsive"],
])

H("1.3.2. Công cụ sử dụng", 3)
P("Bảng 1.2: Các công cụ phát triển và triển khai", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Công cụ", "Mục đích", "Mô tả"], [
    ["Python 3.14","Ngôn ngữ lập trình Backend","Phiên bản Python mới nhất, hỗ trợ type hints và async/await"],
    ["Node.js 22","Runtime JavaScript","Môi trường chạy Next.js, hỗ trợ ESM và TypeScript"],
    ["Uvicorn","ASGI Server","Chạy ứng dụng FastAPI với hiệu năng cao"],
    ["pip/venv","Quản lý gói Python","Tạo môi trường ảo, quản lý dependencies"],
    ["npm/pnpm","Quản lý gói Node.js","Cài đặt và quản lý dependencies cho Frontend"],
    ["Git/GitHub","Quản lý mã nguồn","Quản lý phiên bản, sao lưu và chia sẻ mã nguồn"],
    ["VS Code","IDE","Trình soạn thảo mã nguồn chính"],
    ["Docker","Containerization","Đóng gói và triển khai ứng dụng (kế hoạch)"],
    ["Postman","Kiểm thử API","Kiểm thử các endpoint REST API"],
])
P("Việc lựa chọn các công nghệ và công cụ trên dựa trên nguyên tắc ưu tiên các giải pháp mã nguồn mở, có cộng đồng hỗ trợ lớn và phù hợp với yêu cầu kỹ thuật của hệ thống. FastAPI được chọn cho Backend nhờ hiệu năng cao, hỗ trợ async/await tự nhiên và tích hợp SSE cho streaming. Next.js 15 được chọn cho Frontend nhờ khả năng Server-Side Rendering (SSR), routing linh hoạt và hỗ trợ TypeScript. ChromaDB được chọn làm vector database nhờ tính đơn giản, nhẹ và tích hợp tốt với Python.")
BR()

print("Chapter 1 done.")

IMG("h1_1_tongquan_baivan.png", "Hình 1.1: Sơ đồ tổng quan bài toán hệ thống tư vấn nội quy UTC")

# ╔══════════════════════════════════════════════════════════════╗
# ║           CHƯƠNG 2: MÔ HÌNH NGÔN NGỮ LỚN                   ║
# ╚══════════════════════════════════════════════════════════════╝

H("CHƯƠNG 2. MÔ HÌNH NGÔN NGỮ LỚN", 1)

H("2.1. Tổng quan mô hình ngôn ngữ lớn", 2)
P("Mô hình ngôn ngữ lớn (Large Language Model - LLM) là các hệ thống trí tuệ nhân tạo được huấn luyện trên lượng dữ liệu văn bản khổng lồ (hàng nghìn tỷ token) với số lượng tham số rất lớn (từ hàng tỷ đến hàng nghìn tỷ). Không giống như các mô hình NLP truyền thống được thiết kế cho từng tác vụ cụ thể, LLM là các \"foundation models\" – mô hình nền tảng có khả năng tổng quát hóa cho nhiều tác vụ khác nhau mà không cần hoặc cần rất ít fine-tuning.")
P("Lịch sử phát triển của LLM có thể được chia thành ba giai đoạn chính. Giai đoạn 1 (2017-2019) đánh dấu sự ra đời của kiến trúc Transformer (Vaswani et al., 2017) [1] và các mô hình tiên phong như BERT (Devlin et al., 2019) [2] và GPT-1 (Radford et al., 2018) [7]. Giai đoạn 2 (2020-2022) chứng kiến sự bùng nổ về quy mô với GPT-3 175B (Brown et al., 2020) [3] và LLaMA 2 (Touvron et al., 2023) [8], khẳng định rằng các mô hình lớn hơn có khả năng \"học trong ngữ cảnh\" (in-context learning) – thực hiện tác vụ mới chỉ từ vài ví dụ trong prompt mà không cần fine-tuning. Giai đoạn 3 (2023-nay) là kỷ nguyên của các LLM thương mại và mã nguồn mở quy mô lớn như GPT-4, Claude 3.5, Gemini 1.5, LLaMA 3 và Qwen 3.5, với khả năng xử lý ngữ cảnh lên đến hàng trăm nghìn token và thực hiện đa dạng các tác vụ phức tạp.")
P("Bảng 2.1: Một số mô hình LLM tiêu biểu", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Mô hình", "Tổ chức", "Tham số", "Ngữ cảnh tối đa", "Năm", "Mã nguồn mở"], [
    ["GPT-3","OpenAI","175B","2.048","2020","Không"],
    ["GPT-4","OpenAI","~1.8T*","128K","2023","Không"],
    ["Claude 3.5 Sonnet","Anthropic","—","200K","2024","Không"],
    ["Gemini 1.5 Pro","Google","—","1M","2024","Không"],
    ["LLaMA 2","Meta","7B/13B/70B","4K","2023","Có"],
    ["LLaMA 3.1","Meta","8B/70B/405B","128K","2024","Có"],
    ["Qwen 3.5","Alibaba","—","128K","2025","Có"],
    ["DeepSeek-V3","DeepSeek","671B (MoE)","128K","2024","Có"],
    ["Mistral Large","Mistral","—","128K","2024","Có"],
    ["Gemma 3","Google","1B-27B","128K","2025","Có"],
])
P("(*): Số liệu ước tính, không được công bố chính thức. MoE = Mixture of Experts.")
P("Các LLM hiện đại chia sẻ một số đặc điểm chung: (1) Sử dụng kiến trúc Decoder-only Transformer với mặt nạ nhân quả (causal mask), cho phép sinh văn bản tuần tự (autoregressive); (2) Được huấn luyện qua ba giai đoạn: Pre-training trên dữ liệu lớn → Supervised Fine-Tuning (SFT) trên dữ liệu hội thoại → Alignment (RLHF/DPO) để căn chỉnh với giá trị con người; (3) Hỗ trợ xử lý ngữ cảnh dài thông qua các kỹ thuật như Rotary Position Embedding (RoPE), FlashAttention (Dao et al., 2022) [15] và Grouped Query Attention (GQA). Transformer-XL (Dai et al., 2019) [14] là một trong những mô hình tiên phong mở rộng khả năng xử lý chuỗi dài vượt qua giới hạn ngữ cảnh cố định.")

H("2.2. Kiến trúc Transformer", 2)

H("2.2.1. Tổng quan kiến trúc", 3)
P("Kiến trúc Transformer được giới thiệu bởi Vaswani et al. (2017) trong bài báo \"Attention Is All You Need\" [1] đã tạo ra một cuộc cách mạng trong lĩnh vực xử lý ngôn ngữ tự nhiên. Khác với các kiến trúc trước đó như RNN và LSTM xử lý tuần tự từng token, Transformer xử lý toàn bộ chuỗi đầu vào song song thông qua cơ chế self-attention, cho phép mô hình nắm bắt được các mối quan hệ phụ thuộc xa trong văn bản một cách hiệu quả.")
P("Kiến trúc Transformer gốc được thiết kế cho bài toán dịch máy, bao gồm hai thành phần chính: Encoder (bộ mã hóa) và Decoder (bộ giải mã). Encoder nhận chuỗi đầu vào và tạo ra các biểu diễn ngữ cảnh (contextual representations). Decoder nhận biểu diễn từ Encoder và sinh chuỗi đầu ra theo từng bước. Cả Encoder và Decoder đều được xây dựng từ nhiều tầng (layer) xếp chồng lên nhau, mỗi tầng bao gồm hai sub-layer chính: Multi-Head Self-Attention và Feed-Forward Network, kèm theo kết nối residual và layer normalization.")
P("Ba đặc điểm thiết kế quan trọng của Transformer là: (1) Self-Attention – cho phép mỗi vị trí trong chuỗi \"chú ý\" đến tất cả các vị trí khác để tính toán biểu diễn; (2) Position Encoding – bổ sung thông tin về vị trí của token trong chuỗi, vì self-attention không có khái niệm về thứ tự; (3) Multi-Head Attention – cho phép mô hình học nhiều kiểu quan hệ khác nhau giữa các token cùng một lúc.")

IMG("h2_1_transformer.png", "Hình 2.1: Kiến trúc Transformer (Vaswani et al., 2017) [1]")

H("2.2.2. Encoder – Decoder", 3)
P("Encoder bao gồm N = 6 tầng giống hệt nhau xếp chồng. Mỗi tầng có hai sub-layer: (1) Multi-Head Self-Attention – cho phép mỗi token trong chuỗi đầu vào \"chú ý\" đến tất cả các token khác; (2) Position-wise Feed-Forward Network – một mạng neural truyền thẳng (fully connected) áp dụng độc lập cho từng vị trí. Mỗi sub-layer được bao quanh bởi kết nối residual (cộng đầu vào với đầu ra của sub-layer) và theo sau bởi layer normalization. Công thức tổng quát cho mỗi sub-layer là: LayerNorm(x + Sublayer(x)). Mô hình T5 (Raffel et al., 2020) [11] là đại diện tiêu biểu cho kiến trúc Encoder-Decoder đầy đủ, với triết lý \"mọi tác vụ NLP đều có thể được chuyển thành bài toán text-to-text\".")
P("Decoder cũng bao gồm N = 6 tầng, nhưng mỗi tầng có ba sub-layer: (1) Masked Multi-Head Self-Attention – tương tự như trong Encoder nhưng có thêm mặt nạ (mask) để ngăn mỗi vị trí \"nhìn thấy\" các vị trí tương lai trong chuỗi, đảm bảo tính tự hồi quy (autoregressive) khi sinh văn bản; (2) Cross-Attention – cho phép Decoder \"chú ý\" đến đầu ra của Encoder, kết nối thông tin giữa đầu vào và đầu ra; (3) Feed-Forward Network – giống như trong Encoder.")
P("Trong thực tế, các LLM hiện đại hầu hết chỉ sử dụng thành phần Decoder (Decoder-only Transformer), bỏ qua Encoder. Các mô hình như GPT, LLaMA và Qwen đều thuộc họ này. Ưu điểm của Decoder-only là tính đơn giản và khả năng sinh văn bản tự nhiên. Ngược lại, các mô hình Encoder-only như BERT phù hợp cho các tác vụ hiểu văn bản (phân loại, NER, QA) nhưng không thể sinh văn bản. Các biến thể như RoBERTa (Liu et al., 2019) [12] đã cải thiện đáng kể hiệu năng của BERT thông qua việc tối ưu hóa quy trình huấn luyện với nhiều dữ liệu hơn và loại bỏ tác vụ NSP.")

H("2.2.3. Cơ chế self-attention", 3)
P("Self-attention là trái tim của kiến trúc Transformer. Cơ chế này cho phép mô hình tính toán mức độ \"liên quan\" giữa mọi cặp vị trí trong chuỗi đầu vào, từ đó tạo ra biểu diễn ngữ cảnh cho từng vị trí. Cụ thể, với mỗi token trong chuỗi, self-attention tính toán ba vector: Query (Q), Key (K) và Value (V) thông qua các phép nhân ma trận với các ma trận trọng số học được.")
P("Công thức Scaled Dot-Product Attention:")
P("Attention(Q, K, V) = softmax(QK^T / √d_k) × V", b=True)

IMG("h2_2_attention.png", "Hình 2.2: Cơ chế Scaled Dot-Product Attention")
P("Trong đó: Q (Query) đại diện cho câu hỏi \"token này đang tìm kiếm thông tin gì?\"; K (Key) đại diện cho \"token này chứa thông tin gì?\"; V (Value) là thông tin thực tế được truyền đi; d_k là số chiều của vector Key, được dùng để chia tỉ lệ (scale) nhằm tránh giá trị dot product quá lớn gây ra gradient vanishing sau softmax.")
P("Multi-Head Attention mở rộng ý tưởng này bằng cách chạy nhiều cơ chế self-attention song song (mỗi \"head\" học một kiểu quan hệ khác nhau), sau đó nối kết quả lại và chiếu về không gian ban đầu. Công thức:")
P("MultiHead(Q, K, V) = Concat(head_1, ..., head_h) × W^O", b=True)
P("Mỗi head_i = Attention(Q × W_i^Q, K × W_i^K, V × W_i^V). Transformer gốc sử dụng h = 8 head, mỗi head có d_k = d_v = d_model/h = 64. Việc sử dụng nhiều head cho phép mô hình học được nhiều loại quan hệ khác nhau giữa các token: quan hệ cú pháp (chủ ngữ - động từ), quan hệ ngữ nghĩa (đồng nghĩa, liên quan), quan hệ vị trí (xa - gần), và nhiều loại quan hệ khác.")

IMG("h2_3_multihead.png", "Hình 2.3: Cơ chế Multi-Head Attention")

H("2.3. Prompt Engineering", 2)

IMG("h2_4_prompt_engineering.png", "Hình 2.4: Sơ đồ minh họa các chiến lược Prompt Engineering")
P("Prompt Engineering là nghệ thuật và khoa học thiết kế đầu vào (prompt) cho LLM để đạt được đầu ra mong muốn. Đây là một kỹ năng quan trọng khi làm việc với LLM, vì chất lượng prompt ảnh hưởng trực tiếp đến chất lượng câu trả lời. Không giống như fine-tuning yêu cầu huấn luyện lại mô hình, Prompt Engineering tận dụng khả năng in-context learning của LLM – khả năng \"học\" từ chính prompt mà không cần cập nhật tham số.")
P("Các chiến lược Prompt Engineering chính bao gồm: Zero-shot Prompting – đưa ra câu hỏi trực tiếp không kèm ví dụ, phù hợp với các tác vụ đơn giản; Few-shot Prompting – cung cấp một vài ví dụ về định dạng đầu vào - đầu ra mong muốn trong prompt, giúp mô hình hiểu được pattern cần tuân theo; Chain-of-Thought (CoT) (Wei et al., 2022) [13] – yêu cầu mô hình \"suy nghĩ từng bước\" trước khi đưa ra câu trả lời cuối cùng, đặc biệt hiệu quả cho các bài toán suy luận phức tạp; Self-Consistency – chạy CoT nhiều lần và chọn câu trả lời xuất hiện nhiều nhất, cải thiện độ chính xác; và Tree-of-Thoughts (ToT) – cho phép mô hình khám phá nhiều nhánh suy luận và quay lui khi cần.")
P("Bảng 2.2: So sánh các chiến lược Prompt Engineering", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Chiến lược", "Mô tả", "Ưu điểm", "Nhược điểm", "Ứng dụng"], [
    ["Zero-shot","Không cung cấp ví dụ","Đơn giản, nhanh","Độ chính xác thấp với tác vụ phức tạp","Phân loại đơn giản, dịch cơ bản"],
    ["Few-shot","Cung cấp 2-5 ví dụ","Cải thiện đáng kể chất lượng","Tốn token, cần chọn ví dụ tốt","Phân loại, trích xuất, định dạng"],
    ["Chain-of-Thought","Yêu cầu suy luận từng bước","Mạnh cho toán, logic","Tốn token, có thể lan man","Giải toán, suy luận đa bước"],
    ["Self-Consistency","Chạy CoT nhiều lần, bầu chọn","Tăng độ chính xác","Tốn nhiều token và thời gian","Bài toán cần độ chính xác cao"],
    ["Tree-of-Thoughts","Khám phá nhiều nhánh suy luận","Linh hoạt, giải được bài khó","Rất tốn token, phức tạp","Bài toán tối ưu, lập kế hoạch"],
])
P("Trong hệ thống UTC Assistant, Prompt Engineering đóng vai trò then chốt trong việc đảm bảo chất lượng câu trả lời. System prompt dài 595 dòng được thiết kế cẩn thận bằng tiếng Việt, bao gồm: (1) Định nghĩa vai trò – \"Bạn là trợ lý tư vấn nội quy của Trường Đại học Giao thông Vận tải\"; (2) Quy tắc trả lời – chỉ trả lời dựa trên ngữ cảnh được cung cấp, không suy đoán; (3) Định dạng đầu ra – sử dụng cấu trúc rõ ràng, dễ đọc, không sử dụng Markdown; (4) Xử lý trường hợp ngoài phạm vi – khi không tìm thấy thông tin, trả lời lịch sự và đề xuất kênh hỗ trợ khác. Kỹ thuật Chain-of-Thought cũng được áp dụng ngầm thông qua việc yêu cầu LLM trích dẫn nguồn cụ thể cho từng thông tin trong câu trả lời.")
BR()

print("Chapter 2 done.")

# ╔══════════════════════════════════════════════════════════════╗
# ║        CHƯƠNG 3: TRUY VẤN TẠO SINH TĂNG CƯỜNG (RAG)        ║
# ╚══════════════════════════════════════════════════════════════╝

H("CHƯƠNG 3. TRUY VẤN TẠO SINH TĂNG CƯỜNG", 1)

H("3.1. Tổng quan về RAG", 2)
P("Truy vấn tạo sinh tăng cường (Retrieval-Augmented Generation - RAG) là một kỹ thuật kết hợp giữa truy xuất thông tin (Information Retrieval - IR) và sinh văn bản (Text Generation) để tạo ra các câu trả lời chính xác, có căn cứ dựa trên tài liệu tham khảo. Kỹ thuật này được Lewis et al. (2020) [4] giới thiệu lần đầu tiên tại hội nghị NeurIPS 2020, và nhanh chóng trở thành một trong những phương pháp được sử dụng rộng rãi nhất để xây dựng các hệ thống hỏi đáp dựa trên LLM.")
P("Ý tưởng cốt lõi của RAG rất trực quan: thay vì yêu cầu LLM trả lời câu hỏi chỉ dựa trên kiến thức đã được huấn luyện (có thể lỗi thời, không đầy đủ hoặc không chính xác), RAG trước tiên truy xuất (retrieve) các đoạn văn bản liên quan từ một cơ sở tri thức bên ngoài, sau đó cung cấp các đoạn văn bản này như ngữ cảnh (context) cho LLM để sinh câu trả lời. Cách tiếp cận này mang lại ba lợi ích chính:")
P("(1) Cập nhật tri thức dễ dàng: Kiến thức của hệ thống nằm trong cơ sở dữ liệu vector, có thể được cập nhật, bổ sung hoặc chỉnh sửa mà không cần huấn luyện lại LLM – một quá trình tốn kém và mất thời gian. Khi nhà trường ban hành quy định mới, chỉ cần thêm tài liệu vào cơ sở tri thức là hệ thống có thể trả lời các câu hỏi liên quan.")
P("(2) Giảm thiểu ảo giác (hallucination): Bằng cách neo câu trả lời vào các đoạn văn bản cụ thể được truy xuất, RAG hạn chế đáng kể xu hướng \"bịa đặt\" thông tin của LLM. Trong bối cảnh tư vấn nội quy, đây là yêu cầu bắt buộc – thông tin sai lệch về học phí, điều kiện tốt nghiệp hay quy định kỷ luật có thể gây hậu quả nghiêm trọng.")
P("(3) Minh bạch và có thể kiểm chứng: RAG cho phép trích dẫn nguồn cụ thể cho từng thông tin trong câu trả lời. Người dùng có thể kiểm tra lại thông tin từ tài liệu gốc, tăng độ tin cậy của hệ thống. Hệ thống UTC Assistant hiển thị rõ đoạn văn bản nguồn (có trích dẫn phần, chương, trang) kèm theo mỗi câu trả lời.")

IMG("h3_1_rag_architecture.png", "Hình 3.1: Kiến trúc tổng quan RAG (Lewis et al., 2020) [4]")

H("3.2. Phân loại RAG", 2)
P("Kể từ khi được giới thiệu, RAG đã phát triển thành nhiều biến thể khác nhau, phù hợp với các yêu cầu và mức độ phức tạp khác nhau. Có thể phân loại RAG thành ba nhóm chính dựa trên mức độ tinh chỉnh và tối ưu hóa:")
P("(1) Naive RAG (RAG ngây thơ): Đây là dạng RAG cơ bản nhất, tuân theo quy trình chuẩn: Indexing → Retrieval → Generation. Hệ thống chia tài liệu thành các chunk, tạo embedding, lưu vào vector database. Khi có câu hỏi, hệ thống tạo embedding cho câu hỏi, tìm kiếm các chunk tương tự nhất và đưa chúng vào prompt của LLM. Naive RAG đơn giản, dễ triển khai nhưng có thể gặp vấn đề về chất lượng truy xuất và độ liên quan của ngữ cảnh.")
P("(2) Advanced RAG (RAG nâng cao): Bổ sung các kỹ thuật tối ưu hóa cho Naive RAG ở cả hai giai đoạn retrieval và generation. Ở giai đoạn retrieval: query expansion (mở rộng câu truy vấn), hybrid search (kết hợp sparse + dense retrieval), re-ranking (sắp xếp lại kết quả). Ở giai đoạn generation: prompt compression (nén ngữ cảnh), context curation (chọn lọc ngữ cảnh).")
P("(3) Modular RAG (RAG mô-đun): Kiến trúc linh hoạt nhất, cho phép tổ hợp các mô-đun khác nhau tùy theo bài toán cụ thể. Các mô-đun có thể bao gồm: memory module (ghi nhớ hội thoại trước đó), routing module (định tuyến câu hỏi đến nguồn dữ liệu phù hợp), fusion module (kết hợp thông tin từ nhiều nguồn).")
P("Bảng 3.1: So sánh các phân loại RAG", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Loại RAG", "Độ phức tạp", "Chất lượng truy xuất", "Chi phí", "Phù hợp khi"], [
    ["Naive RAG","Thấp","Cơ bản","Thấp","Dự án nhỏ, POC, dữ liệu đơn giản"],
    ["Advanced RAG","Trung bình","Cao","Trung bình","Hệ thống sản xuất, yêu cầu chính xác"],
    ["Modular RAG","Cao","Rất cao","Cao","Hệ thống phức tạp, đa nguồn dữ liệu"],
])
P("Hệ thống UTC Assistant sử dụng phương pháp Advanced RAG với pipeline 3 tầng: Bi-encoder retrieval → BM25 hybrid → LLM Reranker. Ngoài ra, hệ thống còn tích hợp query expansion (mở rộng từ đồng nghĩa tiếng Việt), topic filter (lọc theo phần/chương) và MMR (Maximal Marginal Relevance) để tăng tính đa dạng của kết quả truy xuất.")

IMG("h3_2_rag_variants.png", "Hình 3.2: Phân loại các biến thể RAG")

H("3.3. Quy trình thực hiện phương pháp RAG ngây thơ (Naive RAG)", 2)
P("Naive RAG là nền tảng để hiểu các biến thể phức tạp hơn. Quy trình Naive RAG bao gồm bốn bước chính: Indexing (lập chỉ mục), Embedding (nhúng từ), Retrieving (truy xuất) và Generating (sinh câu trả lời). Phần này sẽ mô tả chi tiết từng bước.")

IMG("h3_3_naive_rag.png", "Hình 3.3: Quy trình Naive RAG: Indexing → Retrieval → Generation")

H("3.3.1. Indexing", 3)
P("Indexing là bước đầu tiên và quan trọng nhất trong pipeline RAG, quyết định chất lượng của toàn bộ hệ thống. Mục tiêu của Indexing là chuyển đổi tài liệu thô (PDF, Word, HTML...) thành các đơn vị nhỏ hơn (chunk) có thể tìm kiếm được, kèm theo metadata để hỗ trợ truy xuất.")
P("Quy trình Indexing trong UTC Assistant bao gồm các bước sau:")
P("Bước 1 – Trích xuất văn bản (OCR): File PDF Sổ tay sinh viên được xử lý qua công cụ Marker-PDF để trích xuất văn bản và cấu trúc tài liệu (mục lục, tiêu đề, bảng biểu). Kết quả là một file văn bản thuần túy với cấu trúc được bảo toàn.")
P("Bước 2 – Phân tích cấu trúc (Structured Chunking): Thay vì chia nhỏ tài liệu một cách mù quáng theo số lượng ký tự cố định, UTC Assistant sử dụng phương pháp structured chunking – chia nhỏ tài liệu dựa trên cấu trúc logic: Phần (Part) → Chương (Chapter) → Mục (Section). Mỗi chunk tương ứng với một đơn vị nội dung có ý nghĩa hoàn chỉnh (thường là một mục hoặc một nhóm điều khoản liên quan). Phương pháp này đảm bảo mỗi chunk chứa đủ ngữ cảnh để LLM có thể hiểu và sinh câu trả lời chính xác.")
P("Bước 3 – Tạo metadata: Mỗi chunk được gán metadata phong phú để hỗ trợ truy xuất có chọn lọc. Metadata bao gồm: mã phần (phan_so), tiêu đề phần, tiêu đề chương, số trang trong PDF gốc, và từ khóa chủ đề. Metadata cho phép hệ thống lọc nhanh các chunk theo chủ đề trước khi thực hiện truy xuất ngữ nghĩa, cải thiện đáng kể độ chính xác.")
P("Bước 4 – Lưu trữ vào Vector Database: Các chunk được lưu vào ChromaDB – một cơ sở dữ liệu vector mã nguồn mở, nhẹ và dễ sử dụng. ChromaDB lưu trữ cả vector embedding và metadata, hỗ trợ tìm kiếm kết hợp (hybrid search) giữa similarity search và metadata filtering.")

H("3.3.2. Nhúng từ (Embedding)", 3)
P("Embedding là quá trình chuyển đổi văn bản thành các vector số trong không gian nhiều chiều, sao cho các văn bản có ý nghĩa tương tự sẽ có vector gần nhau trong không gian này. Đây là cầu nối quan trọng giữa ngôn ngữ tự nhiên và tính toán số học, cho phép máy tính \"hiểu\" và so sánh ngữ nghĩa của văn bản.")
P("Trong pipeline RAG, embedding được sử dụng ở hai thời điểm: (1) Indexing time – tạo embedding cho tất cả các chunk trong cơ sở tri thức và lưu vào vector database; (2) Query time – tạo embedding cho câu hỏi của người dùng để tìm kiếm các chunk tương tự.")
P("Bảng 3.2: Các phương pháp embedding tiêu biểu", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Phương pháp", "Mô hình đại diện", "Số chiều", "Đặc điểm"], [
    ["Sparse (TF-IDF, BM25)","—","|V| (từ điển)","Dựa trên tần suất từ, nhanh, không hiểu ngữ nghĩa"],
    ["Static Embedding","Word2Vec, GloVe","100-300","Vector cố định cho mỗi từ, không ngữ cảnh"],
    ["Contextual (Encoder)","BERT, RoBERTa","768-1024","Vector phụ thuộc ngữ cảnh, hiểu nghĩa từ trong câu"],
    ["Sentence Embedding","SBERT (Reimers & Gurevych, 2019) [10], E5, BGE","768-4096","Vector cho cả câu/đoạn, tối ưu cho similarity search"],
    ["Multilingual","BAAI/bge-m3","1024","Đa ngôn ngữ, hỗ trợ cả sparse + dense retrieval"],
])
P("Hệ thống UTC Assistant sử dụng mô hình BAAI/bge-m3 (Chen et al., 2024) [6] – một mô hình embedding đa ngôn ngữ tiên tiến dựa trên XLM-RoBERTa. bge-m3 tạo ra vector 1024 chiều cho mỗi đoạn văn bản và hỗ trợ cả dense retrieval (tìm kiếm dựa trên vector dày đặc) và sparse retrieval (tìm kiếm dựa trên trọng số từ). Độ tương đồng giữa hai vector embedding được đo bằng cosine similarity: sim(u, v) = (u·v) / (||u|| × ||v||). Hai văn bản càng giống nhau về ngữ nghĩa thì cosine similarity càng gần 1.")

IMG("h3_4_embedding.png", "Hình 3.4: Minh họa quá trình embedding văn bản thành vector")

H("3.3.3. Retrieving", 3)
P("Retrieving là quá trình tìm kiếm và truy xuất các chunk liên quan nhất từ cơ sở tri thức dựa trên câu hỏi của người dùng. Đây là bước then chốt quyết định chất lượng của câu trả lời cuối cùng – nếu truy xuất sai chunk, LLM sẽ không có đủ thông tin để trả lời chính xác dù có mạnh đến đâu.")
P("Các chiến lược retrieval chính:")
P("(1) Sparse Retrieval (Truy xuất thưa): Dựa trên tần suất từ và thống kê văn bản, tiêu biểu là BM25 (Robertson & Zaragoza, 2009) [9]. Ưu điểm: nhanh, hiệu quả với truy vấn chứa từ khóa cụ thể. Nhược điểm: không hiểu được ngữ nghĩa, từ đồng nghĩa, hoặc paraphrase.")
P("(2) Dense Retrieval (Truy xuất dày đặc): Dựa trên embedding vector và similarity search. Ưu điểm: hiểu ngữ nghĩa, bắt được paraphrase và từ đồng nghĩa. Nhược điểm: chậm hơn sparse, có thể bỏ sót từ khóa chính xác.")
P("(3) Hybrid Retrieval (Truy xuất kết hợp): Kết hợp cả sparse và dense retrieval, sau đó dùng Reciprocal Rank Fusion (RRF) để hợp nhất kết quả. Đây là phương pháp được UTC Assistant sử dụng, kết hợp BM25 (sparse) và Bi-encoder (dense) để tận dụng ưu điểm của cả hai.")
P("(4) Reranking: Sau khi có danh sách candidate chunks từ hybrid retrieval, một mô hình reranker (có thể là LLM hoặc cross-encoder) sẽ chấm điểm lại và sắp xếp các chunk theo độ liên quan thực sự với câu hỏi. UTC Assistant sử dụng LLM Reranker – gọi LLM để đánh giá mức độ liên quan của từng chunk với câu hỏi, sau đó chọn top-K chunk liên quan nhất.")
P("Bảng 3.3: So sánh các chiến lược retrieval", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Chiến lược", "Cơ chế", "Ưu điểm", "Nhược điểm", "Sử dụng trong UTC Assistant"], [
    ["Sparse (BM25)","Tần suất từ, TF-IDF","Nhanh, chính xác với từ khóa","Không hiểu ngữ nghĩa, paraphrase","Tầng 2: BM25 Hybrid"],
    ["Dense (Bi-encoder)","Embedding vector, cosine similarity","Hiểu ngữ nghĩa, đa ngôn ngữ","Chậm hơn sparse","Tầng 1: Bi-encoder retrieval"],
    ["Hybrid (RRF)","Kết hợp sparse + dense","Tận dụng ưu điểm cả hai","Phức tạp hơn","Hợp nhất kết quả tầng 1+2"],
    ["LLM Reranker","LLM chấm điểm liên quan","Rất chính xác","Tốn token, chậm","Tầng 3: Chọn top-K cuối cùng"],
])
P("Ngoài ra, UTC Assistant còn áp dụng các kỹ thuật tối ưu retrieval: Query Expansion – mở rộng câu truy vấn với từ đồng nghĩa tiếng Việt (ví dụ: \"học phí\" → \"học phí, đóng tiền, nộp tiền, thanh toán\"); Topic Filter – lọc chunk theo metadata (phần, chương) dựa trên từ khóa trong câu hỏi; và MMR (Maximal Marginal Relevance) – đảm bảo các chunk được chọn không trùng lặp về nội dung, tăng tính đa dạng của ngữ cảnh.")

H("3.3.4. Generating", 3)
P("Generating là bước cuối cùng trong pipeline RAG, nơi LLM sinh ra câu trả lời dựa trên câu hỏi của người dùng và ngữ cảnh được truy xuất từ bước Retrieving. Chất lượng của bước này phụ thuộc vào ba yếu tố: (1) Chất lượng của LLM được sử dụng; (2) Chất lượng và mức độ liên quan của ngữ cảnh được cung cấp; (3) Cách thiết kế prompt để hướng dẫn LLM sử dụng ngữ cảnh hiệu quả.")
P("Quy trình Generating trong UTC Assistant:")
P("Bước 1 – Xây dựng prompt: System prompt (595 dòng tiếng Việt) được thiết kế để định nghĩa vai trò, quy tắc và định dạng đầu ra. User prompt bao gồm: (a) Ngữ cảnh – các chunk được truy xuất, kèm theo trích dẫn nguồn (phần, chương, trang); (b) Lịch sử hội thoại – để hỗ trợ multi-turn (hội thoại nhiều lượt); (c) Câu hỏi hiện tại của người dùng.")
P("Bước 2 – Gọi LLM: Hệ thống gửi prompt đã xây dựng đến LLM qwen35-opus thông qua API OpenAI-compatible endpoint (endpoint nội bộ tại 100.64.0.25:8103). LLM được cấu hình với temperature = 0.1 (thấp để đảm bảo tính nhất quán), max_tokens = 3000, và streaming = True để phản hồi theo thời gian thực.")
P("Bước 3 – Streaming phản hồi: Câu trả lời được stream về client qua Server-Sent Events (SSE), cho phép hiển thị từng token ngay khi được sinh ra, tạo trải nghiệm tương tác mượt mà cho người dùng. Đồng thời, metadata về nguồn tham khảo được gửi kèm để hiển thị trích dẫn.")
P("Bước 4 – Hậu xử lý: Câu trả lời hoàn chỉnh được lưu vào cơ sở dữ liệu cùng với thông tin về câu hỏi, ngữ cảnh đã sử dụng, thời gian phản hồi và đánh giá của người dùng (nếu có). Dữ liệu này được sử dụng cho mục đích thống kê và cải thiện hệ thống.")
P("Một điểm đáng chú ý trong khâu Generating của UTC Assistant là cơ chế graceful degradation: nếu không tìm thấy chunk nào liên quan (dưới ngưỡng similarity), hệ thống sẽ thông báo lịch sự rằng câu hỏi nằm ngoài phạm vi kiến thức và đề xuất các kênh hỗ trợ khác (Phòng Công tác Chính trị - Sinh viên, cố vấn học tập), thay vì cố gắng sinh câu trả lời dựa trên kiến thức tổng quát của LLM – vốn có thể không chính xác cho các quy định cụ thể của UTC.")
BR()

print("Chapter 3 done.")

# ╔══════════════════════════════════════════════════════════════╗
# ║    CHƯƠNG 4: XÂY DỰNG HỆ THỐNG TƯ VẤN NỘI QUY UTC         ║
# ╚══════════════════════════════════════════════════════════════╝

H("CHƯƠNG 4. XÂY DỰNG HỆ THỐNG TƯ VẤN NỘI QUY TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", 1)

H("4.1. Phân tích thiết kế hệ thống", 2)
P("Trước khi tiến hành xây dựng hệ thống, bước phân tích thiết kế được thực hiện để xác định rõ các yêu cầu chức năng, đối tượng người dùng và kiến trúc tổng thể. Hệ thống UTC Assistant được thiết kế với hai nhóm người dùng chính: Sinh viên (có nhu cầu tra cứu nội quy) và Quản trị viên (có nhu cầu quản lý, giám sát hệ thống).")

H("4.1.1. Sơ đồ ca sử dụng", 3)
P("Sơ đồ ca sử dụng (Use Case Diagram) mô tả các tác nhân (actor) và các chức năng (use case) mà hệ thống cung cấp. Hệ thống UTC Assistant có hai tác nhân chính:")
P("(1) Sinh viên (Student): Là người dùng chính của hệ thống, có nhu cầu tra cứu thông tin về nội quy, quy chế nhà trường. Sinh viên có thể tương tác với chatbot để đặt câu hỏi bằng ngôn ngữ tự nhiên, nhận câu trả lời kèm trích dẫn nguồn, xem lịch sử hội thoại và đánh giá chất lượng câu trả lời.")
P("(2) Quản trị viên (Admin): Là cán bộ phụ trách quản lý hệ thống, có quyền truy cập vào trang quản trị. Admin có thể đăng nhập (xác thực), xem thống kê về hoạt động của hệ thống (số lượt hỏi, tỉ lệ hài lòng, câu hỏi phổ biến...), cập nhật tài liệu nguồn (upload Sổ tay sinh viên mới, trigger re-index), và quản lý cơ sở tri thức.")
P("Sơ đồ ca sử dụng chi tiết được trình bày trong Hình 4.1, Hình 4.2 và Hình 4.3 của báo cáo. Các ca sử dụng chính bao gồm: Tra cứu nội quy, Đánh giá câu trả lời, Xem lịch sử hội thoại (nhóm Sinh viên); Đăng nhập, Xem thống kê, Cập nhật dữ liệu, Quản lý tài liệu (nhóm Admin).")

IMG("h4_1_usecase_tongquan.png", "Hình 4.1: Sơ đồ ca sử dụng tổng quan hệ thống UTC Assistant")
IMG("h4_2_usecase_sinhvien.png", "Hình 4.2: Sơ đồ ca sử dụng chi tiết – Nhóm chức năng sinh viên")
IMG("h4_3_usecase_admin.png", "Hình 4.3: Sơ đồ ca sử dụng chi tiết – Nhóm chức năng quản trị")

H("4.1.2. Đặc tả các chức năng", 3)
P("Phần này đặc tả chi tiết các chức năng chính của hệ thống dưới dạng bảng, bao gồm: mã chức năng, mô tả, điều kiện tiên quyết, luồng sự kiện chính và kết quả mong đợi.")

P("Bảng 4.1: Đặc tả ca sử dụng – Tra cứu nội quy (UC-01)", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Thuộc tính", "Mô tả"], [
    ["Mã ca sử dụng","UC-01"],
    ["Tên","Tra cứu nội quy"],
    ["Tác nhân","Sinh viên"],
    ["Mô tả","Sinh viên đặt câu hỏi về nội quy và nhận câu trả lời từ chatbot"],
    ["Tiền điều kiện","Hệ thống đã được khởi tạo với dữ liệu Sổ tay sinh viên"],
    ["Luồng chính","1. Sinh viên nhập câu hỏi vào ô chat\n2. Hệ thống embedding câu hỏi\n3. Truy xuất top-K chunk liên quan từ ChromaDB\n4. LLM sinh câu trả lời dựa trên ngữ cảnh\n5. Hiển thị câu trả lời (streaming) + nguồn tham khảo"],
    ["Kết quả","Câu trả lời chính xác, có trích dẫn nguồn cụ thể"],
])

P("Bảng 4.2: Đặc tả ca sử dụng – Quản lý dữ liệu (UC-02)", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Thuộc tính", "Mô tả"], [
    ["Mã ca sử dụng","UC-02"],
    ["Tên","Quản lý dữ liệu"],
    ["Tác nhân","Quản trị viên"],
    ["Mô tả","Admin cập nhật tài liệu nguồn và quản lý cơ sở tri thức"],
    ["Tiền điều kiện","Admin đã đăng nhập thành công"],
    ["Luồng chính","1. Admin upload file PDF Sổ tay sinh viên mới\n2. Hệ thống OCR và trích xuất cấu trúc\n3. Chia chunk và tạo embedding\n4. Cập nhật ChromaDB (add/update/delete)\n5. Xác nhận hoàn tất"],
    ["Kết quả","Cơ sở tri thức được cập nhật với dữ liệu mới"],
])

P("Bảng 4.3: Đặc tả ca sử dụng – Thống kê và báo cáo (UC-03)", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Thuộc tính", "Mô tả"], [
    ["Mã ca sử dụng","UC-03"],
    ["Tên","Thống kê và báo cáo"],
    ["Tác nhân","Quản trị viên"],
    ["Mô tả","Admin xem các chỉ số thống kê về hoạt động của hệ thống"],
    ["Tiền điều kiện","Admin đã đăng nhập thành công"],
    ["Luồng chính","1. Admin truy cập Dashboard\n2. Hệ thống hiển thị: tổng số lượt hỏi, câu hỏi phổ biến, tỉ lệ hài lòng, thời gian phản hồi trung bình\n3. Admin có thể lọc theo khoảng thời gian"],
    ["Kết quả","Báo cáo thống kê trực quan, hỗ trợ ra quyết định"],
])

P("Bảng 4.7: Các quy tắc nghiệp vụ chính trong hệ thống", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Rule ID", "Mô tả", "Điều kiện kích hoạt", "Cảnh báo"], [
    ["R1","Chỉ trả lời dựa trên ngữ cảnh được truy xuất","Tất cả câu hỏi","Không suy đoán nếu không có chunk liên quan"],
    ["R2","Hiển thị trích dẫn nguồn","Mọi câu trả lời","Bắt buộc kèm phần/chương/trang"],
    ["R3","Từ chối câu hỏi ngoài phạm vi","Similarity < ngưỡng","Đề xuất kênh hỗ trợ thay thế"],
    ["R4","Streaming phản hồi theo thời gian thực","Tất cả câu hỏi","Sử dụng SSE, không đợi phản hồi hoàn chỉnh"],
    ["R5","Xác thực Admin trước khi truy cập","Truy cập /admin","Token-based authentication"],
    ["R6","Không lưu thông tin cá nhân","Tất cả tương tác","Không yêu cầu MSSV hoặc định danh"],
    ["R7","Phân lớp câu hỏi trước khi truy xuất","Câu hỏi có từ khóa chủ đề","Lọc chunk theo metadata trước similarity search"],
])

H("4.2. Xây dựng trang web tư vấn viên nội quy", 2)

H("4.2.1. Cài đặt chatbot sử dụng phương pháp RAG", 3)
P("Trang web tư vấn viên nội quy là giao diện chính dành cho sinh viên, cung cấp chatbot hỏi đáp về nội quy trường. Phần này mô tả chi tiết kiến trúc và cách cài đặt chatbot sử dụng phương pháp RAG.")
P("Kiến trúc tổng thể của hệ thống được thiết kế theo mô hình client-server với ba tầng chính:")
P("(1) Tầng Frontend (Next.js 15): Chịu trách nhiệm hiển thị giao diện chatbot, xử lý input từ người dùng, hiển thị streaming response qua SSE. Giao diện được thiết kế với Tailwind CSS, hỗ trợ responsive trên cả desktop và mobile. Các thành phần chính bao gồm: khung chat, ô nhập liệu, danh sách câu hỏi gợi ý, hiển thị nguồn tham khảo và nút đánh giá.")
P("(2) Tầng Backend (FastAPI, Python 3.14): Xử lý toàn bộ logic nghiệp vụ của hệ thống, bao gồm: nhận câu hỏi từ Frontend, thực hiện pipeline RAG (embedding → retrieval → reranking → generation), streaming phản hồi qua SSE, lưu trữ hoạt động vào SQLite và cung cấp API cho trang quản trị. Backend được tổ chức thành các module độc lập: rag_pipeline.py (pipeline RAG chính), structured_chunker.py (xử lý chunking có cấu trúc), và main.py (FastAPI app với các route).")
P("(3) Tầng Dữ liệu (ChromaDB + SQLite): ChromaDB lưu trữ vector embedding và metadata của các chunk, phục vụ truy xuất ngữ nghĩa. SQLite lưu trữ dữ liệu hoạt động (câu hỏi, câu trả lời, đánh giá), cấu hình hệ thống và thông tin người dùng admin.")

IMG("h4_4_kientruc_tongthe.png", "Hình 4.4: Kiến trúc tổng thể hệ thống UTC Assistant")
IMG("h4_5_pipeline_rag.png", "Hình 4.5: Pipeline RAG 3 tầng: Bi-encoder → BM25 → LLM Reranker")
IMG("h4_6_luong_dulieu.png", "Hình 4.6: Luồng xử lý dữ liệu: PDF → OCR → Chunk → ChromaDB")
P("Bảng 4.4: Cấu trúc metadata trong mỗi chunk của ChromaDB", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Trường", "Kiểu", "Mô tả", "Ví dụ"], [
    ["phan_so","string","Số thứ tự phần trong Sổ tay","\"1\", \"2.1\", \"3\""],
    ["phan_title","string","Tiêu đề phần","\"QUY ĐỊNH VỀ ĐÀO TẠO\""],
    ["chuong_title","string","Tiêu đề chương","\"I. ĐĂNG KÝ MÔN HỌC\""],
    ["page","int","Số trang trong PDF gốc","15"],
    ["chunk_id","string","Mã định danh duy nhất của chunk","\"P2.1_ChI_01\""],
    ["topic","string","Chủ đề chính","\"đào tạo\", \"học phí\", \"kỷ luật\""],
])
P("Bảng 4.5: Các API endpoint chính của UTC Assistant", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Endpoint", "Method", "Mô tả", "Auth"], [
    ["/api/chat","POST","Gửi câu hỏi, nhận streaming response (SSE)","Không"],
    ["/api/chat/history","GET","Lấy lịch sử hội thoại theo session","Không"],
    ["/api/chat/feedback","POST","Gửi đánh giá câu trả lời (thumbs up/down)","Không"],
    ["/api/suggestions","GET","Lấy danh sách câu hỏi gợi ý","Không"],
    ["/api/admin/login","POST","Đăng nhập admin, trả về token","Không"],
    ["/api/admin/stats","GET","Lấy dữ liệu thống kê","Token"],
    ["/api/admin/documents","POST","Upload tài liệu mới, trigger re-index","Token"],
    ["/api/admin/documents","GET","Danh sách tài liệu hiện có","Token"],
    ["/api/admin/questions","GET","Danh sách câu hỏi + câu trả lời (phân trang)","Token"],
    ["/api/health","GET","Kiểm tra trạng thái hệ thống","Không"],
])

H("4.2.2. Phân lớp câu hỏi của người dùng", 3)
P("Để tối ưu hiệu quả truy xuất, UTC Assistant thực hiện phân lớp câu hỏi trước khi đưa vào pipeline RAG. Mục tiêu của phân lớp là xác định chủ đề chính của câu hỏi để áp dụng topic filter – lọc chunk theo metadata trước khi thực hiện similarity search, giảm không gian tìm kiếm và tăng độ chính xác.")
P("Bảng 4.6: Các lớp câu hỏi và chiến lược xử lý", b=True, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)
TBL(["Lớp câu hỏi", "Từ khóa nhận diện", "Topic Filter (metadata)", "Chiến lược"], [
    ["Đào tạo","đăng ký môn, tín chỉ, học phần, lịch học","phan_so: \"2.1\"","Truy xuất trong Phần 2.1 – Quy định đào tạo"],
    ["Học phí","học phí, đóng tiền, nộp tiền, thanh toán","phan_so: \"3\" + title_like: \"XIII\"","Truy xuất trong Phần 3, Chương XIII"],
    ["Ký túc xá","ktx, nội trú, phòng ở, ký túc","phan_so: \"3\" + title_like: \"X. Ký túc\"","Truy xuất trong Phần 3, Chương X"],
    ["Kỷ luật","kỷ luật, vi phạm, cảnh cáo, đình chỉ","phan_so: \"2.4\"","Truy xuất trong Phần 2.4 – Kỷ luật"],
    ["Tốt nghiệp","tốt nghiệp, ra trường, bằng, đồ án","phan_so: \"2.1\"","Truy xuất trong Phần 2.1"],
    ["Thủ tục","thủ tục, giấy tờ, hồ sơ, xác nhận","phan_so: \"3\"","Truy xuất trong Phần 3"],
    ["Bảo hiểm","bảo hiểm, y tế, khám bệnh","phan_so: \"3\" + title_like: \"XII\"","Truy xuất trong Phần 3, Chương XII"],
    ["Email/LMS","email, lms, tài khoản, @utc","phan_so: \"3\" + title_like: \"XIV\"","Truy xuất trong Phần 3, Chương XIV"],
    ["Khác/Chung","(không khớp lớp nào)","Không filter","Truy xuất toàn bộ cơ sở tri thức"],
])
P("Phân lớp được thực hiện thông qua việc so khớp từ khóa trong câu hỏi với từ điển TOPIC_FILTERS được định nghĩa trong rag_pipeline.py. Ngoài ra, Query Expansion được áp dụng để mở rộng câu truy vấn với các từ đồng nghĩa tiếng Việt (ví dụ: \"học phí\" → [\"học phí\", \"hoc phi\", \"đóng tiền\", \"nộp tiền\", \"thanh toán\", \"tài chính\"]), đảm bảo truy xuất được các chunk sử dụng cách diễn đạt khác nhau cho cùng một khái niệm.")

H("4.2.3. Giao diện trang web tư vấn viên nội quy", 3)
P("Giao diện trang web tư vấn viên nội quy được thiết kế tối giản, tập trung vào trải nghiệm người dùng. Các nguyên tắc thiết kế chính bao gồm: (1) Đơn giản – giao diện sạch, không gây phân tâm, tập trung vào khung chat; (2) Tốc độ – streaming phản hồi theo thời gian thực, không bắt người dùng chờ đợi; (3) Tin cậy – hiển thị rõ nguồn tham khảo cho từng câu trả lời, kèm trích dẫn phần/chương/trang; (4) Thân thiện – giao diện tiếng Việt hoàn toàn, font chữ dễ đọc, màu sắc hài hòa.")
P("Các thành phần chính của giao diện chatbot:")
P("(1) Thanh tiêu đề (Header): Hiển thị logo UTC, tên hệ thống \"UTC Assistant – Tư vấn Nội quy\" và nút chuyển đổi giữa chế độ sáng/tối (dark mode).")
P("(2) Khung chat (Chat Container): Khu vực chính hiển thị lịch sử hội thoại. Tin nhắn của người dùng hiển thị bên phải (màu xanh UTC), câu trả lời của chatbot hiển thị bên trái (màu trắng/xám). Mỗi câu trả lời kèm theo phần \"Nguồn tham khảo\" có thể mở rộng, hiển thị chunk gốc với trích dẫn cụ thể. Streaming response hiển thị từng token ngay khi được sinh ra, tạo cảm giác phản hồi tức thì.")
P("(3) Ô nhập liệu (Input Box): Đặt ở cuối màn hình, có placeholder \"Nhập câu hỏi của bạn...\", nút gửi và hỗ trợ gửi bằng phím Enter. Khi hệ thống đang xử lý, ô nhập liệu bị vô hiệu hóa và hiển thị animation \"đang nhập...\".")
P("(4) Câu hỏi gợi ý (Suggested Questions): Hiển thị 4-6 câu hỏi mẫu phổ biến bên dưới ô nhập liệu, giúp sinh viên mới làm quen với hệ thống. Các câu hỏi được chọn từ danh sách câu hỏi phổ biến nhất trong cơ sở dữ liệu.")
P("(5) Thanh bên (Sidebar): Chứa danh sách các phiên hội thoại gần đây, cho phép sinh viên xem lại lịch sử tra cứu. Mỗi phiên được lưu với session_id duy nhất, tự động khôi phục khi người dùng quay lại.")

H("4.3. Xây dựng trang web quản trị tư vấn nội quy", 2)
P("Trang web quản trị dành cho cán bộ quản lý, cung cấp các chức năng giám sát và quản lý hệ thống. Trang quản trị yêu cầu xác thực để đảm bảo an toàn dữ liệu.")

H("4.3.1. Giao diện chức năng xác thực", 3)
P("Chức năng xác thực (Authentication) là lớp bảo vệ đầu tiên của trang quản trị. Hệ thống sử dụng cơ chế xác thực dựa trên token (token-based authentication):")
P("(1) Trang đăng nhập: Giao diện tối giản với logo UTC, form đăng nhập gồm username và password. Khi admin submit, thông tin được gửi đến endpoint /api/admin/login. Nếu xác thực thành công, server trả về JWT token có thời hạn, được lưu trong localStorage của trình duyệt.")
P("(2) Middleware bảo vệ: Tất cả các route /admin/* đều được bảo vệ bởi middleware kiểm tra token. Nếu token không tồn tại hoặc hết hạn, người dùng bị chuyển hướng về trang đăng nhập. Backend cũng kiểm tra token ở mỗi request đến API admin.")
P("(3) Quản lý phiên: Token có thời hạn 24 giờ, sau đó admin cần đăng nhập lại. Hệ thống hỗ trợ chức năng đăng xuất (logout) – xóa token khỏi localStorage và chuyển hướng về trang đăng nhập.")

H("4.3.2. Giao diện chức năng thống kê", 3)
P("Chức năng thống kê (Dashboard) cung cấp cái nhìn tổng quan về hoạt động của hệ thống, hỗ trợ cán bộ quản lý trong việc đánh giá hiệu quả và phát hiện vấn đề. Dashboard được thiết kế dưới dạng các card thông tin và biểu đồ trực quan.")
P("Các chỉ số thống kê chính:")
P("(1) Tổng quan: Tổng số lượt hỏi, số lượt hỏi hôm nay, số câu hỏi được đánh giá hài lòng, tỉ lệ hài lòng (%). Các chỉ số này được hiển thị dưới dạng card với số liệu lớn, dễ đọc.")
P("(2) Biểu đồ lượt hỏi theo thời gian: Biểu đồ đường (line chart) hiển thị số lượt hỏi theo ngày/tuần/tháng, cho phép admin phát hiện xu hướng và thời điểm cao điểm. Có thể lọc theo khoảng thời gian (7 ngày, 30 ngày, 90 ngày).")
P("(3) Top câu hỏi phổ biến: Bảng xếp hạng 10 câu hỏi được hỏi nhiều nhất, kèm số lượt hỏi và tỉ lệ hài lòng. Dữ liệu này giúp admin hiểu được nhu cầu thông tin chính của sinh viên.")
P("(4) Phân bố chủ đề: Biểu đồ tròn (pie chart) hoặc biểu đồ cột (bar chart) hiển thị phân bố câu hỏi theo chủ đề (đào tạo, học phí, ký túc xá, kỷ luật...), giúp xác định lĩnh vực nào sinh viên quan tâm nhiều nhất.")
P("(5) Thời gian phản hồi: Thống kê thời gian phản hồi trung bình, p50, p95 và p99 (milliseconds). Chỉ số này quan trọng để đánh giá hiệu năng hệ thống và phát hiện các vấn đề về tốc độ.")

H("4.3.3. Giao diện chức năng cập nhật dữ liệu", 3)
P("Chức năng cập nhật dữ liệu cho phép admin quản lý cơ sở tri thức của hệ thống. Khi nhà trường ban hành Sổ tay sinh viên mới hoặc cập nhật quy định, admin có thể upload tài liệu mới để hệ thống tự động xử lý và cập nhật.")
P("Các chức năng chính:")
P("(1) Upload tài liệu: Giao diện kéo-thả (drag-and-drop) hoặc chọn file, hỗ trợ định dạng PDF. Sau khi upload, hệ thống tự động thực hiện pipeline: OCR → Structured Chunking → Embedding → Lưu vào ChromaDB. Tiến trình được hiển thị qua thanh tiến trình (progress bar).")
P("(2) Quản lý tài liệu: Danh sách các tài liệu đã upload, hiển thị tên file, ngày upload, số chunk đã tạo, trạng thái (active/archived). Admin có thể kích hoạt/vô hiệu hóa tài liệu hoặc xóa tài liệu cũ.")
P("(3) Xem trước chunk: Cho phép admin xem nội dung của từng chunk sau khi indexing, bao gồm metadata và nội dung văn bản. Chức năng này hữu ích để kiểm tra chất lượng chunking và phát hiện lỗi.")
P("(4) Re-index: Khi cần cập nhật toàn bộ cơ sở tri thức (ví dụ: thay đổi tham số chunking), admin có thể kích hoạt re-index – xóa collection cũ và tạo mới từ tất cả tài liệu hiện có.")
P("(5) Test truy vấn: Công cụ test nội bộ cho phép admin nhập câu hỏi và xem kết quả retrieval (danh sách chunk được truy xuất kèm similarity score) trước khi đưa vào generation. Chức năng này giúp admin đánh giá chất lượng truy xuất và điều chỉnh tham số nếu cần.")
BR()

print("Chapter 4 done.")

# ═══════════════════════════ KẾT LUẬN VÀ KIẾN NGHỊ ═══════════════════════════
H("KẾT LUẬN VÀ KIẾN NGHỊ", 1)

H("a. Kết quả đạt được", 2)
P("Qua quá trình nghiên cứu và thực hiện đồ án, các kết quả chính đã đạt được bao gồm:")
P("(1) Về mặt lý thuyết: Đã nghiên cứu và trình bày có hệ thống các kiến thức nền tảng về Mô hình ngôn ngữ lớn (LLM), bao gồm kiến trúc Transformer, cơ chế self-attention, các kỹ thuật Prompt Engineering, và đặc biệt là kỹ thuật Truy vấn tạo sinh tăng cường (RAG). Các kiến thức này được tổng hợp từ nhiều nguồn tài liệu uy tín, bao gồm các bài báo khoa học tại các hội nghị hàng đầu và sách chuyên khảo.")
P("(2) Về mặt thực nghiệm: Đã xây dựng thành công hệ thống UTC Assistant – trợ lý ảo tư vấn nội quy cho Trường Đại học Giao thông Vận tải. Hệ thống bao gồm đầy đủ các thành phần: (a) Pipeline RAG 3 tầng với khả năng truy xuất chính xác (Hit Rate 100%, MRR 0.980); (b) Giao diện chatbot thân thiện cho sinh viên, hỗ trợ streaming real-time và hiển thị nguồn tham khảo; (c) Giao diện quản trị với các chức năng xác thực, thống kê và cập nhật dữ liệu.")
P("(3) Về mặt kỹ thuật: Hệ thống được xây dựng trên nền tảng công nghệ hiện đại (FastAPI Python 3.14, Next.js 15, ChromaDB, Tailwind CSS), áp dụng các kỹ thuật tiên tiến như Query Expansion tiếng Việt, Topic Filter, Hybrid Search (sparse + dense), LLM Reranker và MMR. Hệ thống cũng được thiết kế với khả năng graceful degradation – xử lý tình huống không tìm thấy thông tin một cách lịch sự và hữu ích.")

H("b. Hạn chế", 2)
P("Bên cạnh những kết quả đạt được, hệ thống vẫn còn một số hạn chế:")
P("(1) Phạm vi dữ liệu: Hệ thống hiện chỉ hoạt động với dữ liệu Sổ tay sinh viên K66. Khi có Sổ tay cho các khóa mới hoặc các tài liệu quy định bổ sung, cần cập nhật thủ công. Chưa có cơ chế tự động phát hiện và tích hợp tài liệu mới.")
P("(2) Chất lượng tiếng Việt: Mặc dù LLM qwen35-opus có khả năng xử lý tiếng Việt tốt, đôi khi vẫn gặp lỗi về ngữ pháp hoặc cách diễn đạt chưa tự nhiên. Embedding model bge-m3 tuy đa ngôn ngữ nhưng chưa được tinh chỉnh riêng cho tiếng Việt.")
P("(3) Khả năng mở rộng: Hệ thống hiện chạy trên một máy chủ đơn, chưa được thiết kế cho khả năng mở rộng ngang (horizontal scaling) khi số lượng người dùng tăng cao. Embedding và LLM inference sử dụng endpoint nội bộ, phụ thuộc vào tài nguyên của máy chủ đó.")
P("(4) Đánh giá chất lượng: Việc đánh giá chất lượng câu trả lời hiện dựa trên đánh giá chủ quan của người dùng (thumbs up/down) và một tập test 50 câu hỏi. Chưa có cơ chế đánh giá tự động, liên tục và toàn diện hơn.")

H("c. Hướng phát triển", 2)
P("Dựa trên những kết quả đã đạt được và các hạn chế đã xác định, một số hướng phát triển trong tương lai được đề xuất:")
P("(1) Mở rộng cơ sở tri thức: Tích hợp thêm các tài liệu khác của nhà trường như Quy chế đào tạo, Quy định công tác sinh viên, các văn bản hướng dẫn của Phòng Đào tạo... để hệ thống có thể trả lời đa dạng câu hỏi hơn. Xây dựng cơ chế tự động cập nhật khi có văn bản mới.")
P("(2) Cải thiện chất lượng tiếng Việt: Fine-tune embedding model trên dữ liệu tiếng Việt chuyên ngành giáo dục đại học. Thử nghiệm với các LLM khác có khả năng tiếng Việt tốt hơn (như Gemini 2.5 Pro, GPT-4o). Xây dựng bộ test đánh giá chất lượng tiếng Việt chuyên biệt.")
P("(3) Nâng cao khả năng tương tác: Hỗ trợ nhập liệu bằng giọng nói (speech-to-text) và đọc câu trả lời (text-to-speech), đặc biệt hữu ích cho sinh viên khiếm thị. Tích hợp với các nền tảng nhắn tin phổ biến (Zalo, Messenger, Telegram) để sinh viên có thể tra cứu ngay trên điện thoại.")
P("(4) Cải thiện hiệu năng: Áp dụng caching cho các câu hỏi phổ biến để giảm tải cho LLM. Tối ưu hóa kích thước prompt và context window. Triển khai kiến trúc microservices để có thể mở rộng ngang khi cần.")
P("(5) Phân tích dữ liệu nâng cao: Xây dựng hệ thống phân tích xu hướng câu hỏi của sinh viên qua các năm học, giúp nhà trường phát hiện sớm các vấn đề sinh viên quan tâm và có biện pháp truyền thông phù hợp.")
P("(6) Tích hợp Agent: Cho phép chatbot không chỉ trả lời câu hỏi mà còn thực hiện các tác vụ như: tra cứu điểm, đăng ký môn học, nộp đơn xin xác nhận... thông qua việc tích hợp với các hệ thống hiện có của nhà trường (QLDT, LMS).")
BR()

print("Ket luan done.")

# ═══════════════════════════ DANH MỤC TÀI LIỆU THAM KHẢO ═══════════════════════════
H("DANH MỤC TÀI LIỆU THAM KHẢO", 1)
refs = [
    "[1] Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., & Polosukhin, I. (2017). Attention Is All You Need. In Advances in Neural Information Processing Systems 30 (NeurIPS 2017), pp. 5998–6008.",
    "[2] Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. In Proceedings of NAACL-HLT 2019, pp. 4171–4186.",
    "[3] Brown, T. B., Mann, B., Ryder, N., Subbiah, M., Kaplan, J., Dhariwal, P., ... & Amodei, D. (2020). Language Models are Few-Shot Learners. In Advances in Neural Information Processing Systems 33 (NeurIPS 2020), pp. 1877–1901.",
    "[4] Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... & Kiela, D. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. In Advances in Neural Information Processing Systems 33 (NeurIPS 2020), pp. 9459–9474.",
    "[5] Xiao, C., & Zhu, J. (2025). Foundations of Large Language Models. CRC Press.",
    "[6] Chen, J., Xiao, S., Zhang, P., Luo, K., Lian, D., & Liu, Z. (2024). BGE M3-Embedding: Multi-Lingual, Multi-Functionality, Multi-Granularity Text Embeddings Through Self-Knowledge Distillation. arXiv preprint arXiv:2402.03216.",
    "[7] Radford, A., Narasimhan, K., Salimans, T., & Sutskever, I. (2018). Improving Language Understanding by Generative Pre-Training. OpenAI Technical Report.",
    "[8] Touvron, H., Martin, L., Stone, K., Albert, P., Almahairi, A., Babaei, Y., ... & Scialom, T. (2023). LLaMA 2: Open Foundation and Fine-Tuned Chat Models. arXiv preprint arXiv:2307.09288.",
    "[9] Robertson, S., & Zaragoza, H. (2009). The Probabilistic Relevance Framework: BM25 and Beyond. Foundations and Trends in Information Retrieval, 3(4), 333–389.",
    "[10] Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. In Proceedings of EMNLP-IJCNLP 2019, pp. 3982–3992.",
    "[11] Raffel, C., Shazeer, N., Roberts, A., Lee, K., Narang, S., Matena, M., Zhou, Y., Li, W., & Liu, P. J. (2020). Exploring the Limits of Transfer Learning with a Unified Text-to-Text Transformer. Journal of Machine Learning Research, 21(140), 1–67.",
    "[12] Liu, Y., Ott, M., Goyal, N., Du, J., Joshi, M., Chen, D., Levy, O., Lewis, M., Zettlemoyer, L., & Stoyanov, V. (2019). RoBERTa: A Robustly Optimized BERT Pretraining Approach. arXiv preprint arXiv:1907.11692.",
    "[13] Wei, J., Wang, X., Schuurmans, D., Bosma, M., Ichter, B., Xia, F., Chi, E., Le, Q., & Zhou, D. (2022). Chain-of-Thought Prompting Elicits Reasoning in Large Language Models. In Advances in Neural Information Processing Systems 35 (NeurIPS 2022), pp. 24824–24837.",
    "[14] Dai, Z., Yang, Z., Yang, Y., Carbonell, J., Le, Q. V., & Salakhutdinov, R. (2019). Transformer-XL: Attentive Language Models Beyond a Fixed-Length Context. In Proceedings of ACL 2019, pp. 2978–2988.",
    "[15] Dao, T., Fu, D. Y., Ermon, S., Rudra, A., & Ré, C. (2022). FlashAttention: Fast and Memory-Efficient Exact Attention with IO-Awareness. In Advances in Neural Information Processing Systems 35 (NeurIPS 2022), pp. 16344–16359.",
]
for ref in refs:
    P(ref, s=12, indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)

BR()
print("TLTK done.")

# ═══════════════════════════ SAVE ═══════════════════════════
doc.save(OUT)
print(f"\n=== BAO CAO DA DUOC TAO THANH CONG ===")
print(f"File: {OUT}")
print(f"Kich thuoc: {os.path.getsize(OUT):,} bytes")
print(f"So luong paragraphs: {len(doc.paragraphs)}")
print(f"So luong tables: {len(doc.tables)}")
