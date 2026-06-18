#!/usr/bin/env python3
"""Fix font: Times New Roman 13pt for ALL text in the report."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

INPUT = "/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon/Bao_cao_MHNNL_UTC_Assistant.docx"
OUTPUT = INPUT  # overwrite

doc = Document(INPUT)

def fix_run(run, size=None):
    """Force Times New Roman on a run."""
    run.font.name = 'Times New Roman'
    if size:
        run.font.size = Pt(size)
    # Set East-Asian font
    rPr = run._element.get_or_add_rPr()
    # Remove existing rFonts
    for child in list(rPr):
        if child.tag == qn('w:rFonts'):
            rPr.remove(child)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.insert(0, rFonts)

# ── Fix Normal style ──
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(13)
normal.paragraph_format.line_spacing = 1.5
rPr = normal.element.get_or_add_rPr()
for child in list(rPr):
    if child.tag in (qn('w:rFonts'), qn('w:sz'), qn('w:szCs')):
        rPr.remove(child)
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')
rFonts.set(qn('w:eastAsia'), 'Times New Roman')
rPr.insert(0, rFonts)
sz = OxmlElement('w:sz')
sz.set(qn('w:val'), '26')  # 13pt = 26 half-points
rPr.append(sz)

# ── Fix ALL paragraphs ──
count = 0
for p in doc.paragraphs:
    # Check if heading
    is_heading = p.style.name.startswith('Heading') if p.style else False
    
    for run in p.runs:
        if is_heading:
            # Headings: keep bold, font Times New Roman
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
            rPr = run._element.get_or_add_rPr()
            for child in list(rPr):
                if child.tag == qn('w:rFonts'):
                    rPr.remove(child)
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            rPr.insert(0, rFonts)
        else:
            fix_run(run, size=13)
        count += 1

# ── Fix ALL tables ──
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    fix_run(run, size=11)  # Tables: 11pt

# ── Fix heading styles ──
for level in [1, 2, 3]:
    try:
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Times New Roman'
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        rPr = h_style.element.get_or_add_rPr()
        for child in list(rPr):
            if child.tag == qn('w:rFonts'):
                rPr.remove(child)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rPr.insert(0, rFonts)
    except KeyError:
        pass

# ── Set default paragraph font via XML ──
# This ensures new runs use Times New Roman
styles_element = doc.styles.element
doc_defaults = styles_element.find(qn('w:docDefaults'))
if doc_defaults is None:
    doc_defaults = OxmlElement('w:docDefaults')
    styles_element.insert(0, doc_defaults)

rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
if rPrDefault is None:
    rPrDefault = OxmlElement('w:rPrDefault')
    doc_defaults.append(rPrDefault)

rPr_elem = rPrDefault.find(qn('w:rPr'))
if rPr_elem is None:
    rPr_elem = OxmlElement('w:rPr')
    rPrDefault.append(rPr_elem)

# Clear existing defaults
for child in list(rPr_elem):
    rPr_elem.remove(child)

rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')
rFonts.set(qn('w:eastAsia'), 'Times New Roman')
rPr_elem.append(rFonts)

sz = OxmlElement('w:sz')
sz.set(qn('w:val'), '26')  # 13pt
rPr_elem.append(sz)

szCs = OxmlElement('w:szCs')
szCs.set(qn('w:val'), '26')
rPr_elem.append(szCs)

# Save
doc.save(OUTPUT)
print(f"✅ Fixed {count} runs in {INPUT}")
print(f"   Font: Times New Roman 13pt (body), 11pt (tables)")
print(f"   Size: {os.path.getsize(OUTPUT):,} bytes")
