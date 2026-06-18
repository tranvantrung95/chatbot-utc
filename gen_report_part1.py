#!/usr/bin/env python3
"""Generate Báo cáo môn học Mô hình ngôn ngữ lớn - UTC Assistant"""

import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUTPUT_DIR = "/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Bao_cao_MHNNL_UTC_Assistant.docx")

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# ── Style helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Force Times New Roman on East Asian runs too
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')
rPr.insert(0, rFonts)

def add_heading_styled(text, level=1):
    """Add heading with Times New Roman."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rPr.insert(0, rFonts)
    return h

def add_para(text, bold=False, italic=False, alignment=None, font_size=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    # Force font
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)
    return p

def add_table_with_style(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        for p in header_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    
    # Data rows
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            for p in row_cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    doc.add_paragraph()  # spacing
    return table

def add_page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# TRANG BÌA (COVER PAGE)
# ═══════════════════════════════════════════════════════════

# Empty lines for centering
for _ in range(4):
    doc.add_paragraph()

add_para("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
add_para("KHOA CÔNG NGHỆ THÔNG TIN", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)

for _ in range(3):
    doc.add_paragraph()

add_para("BÁO CÁO MÔN HỌC", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=16)
add_para("MÔ HÌNH NGÔN NGỮ LỚN", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)

for _ in range(2):
    doc.add_paragraph()

add_para("ĐỀ TÀI", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)
add_para("XÂY DỰNG HỆ THỐNG TRỢ LÝ ẢO HỖ TRỢ SINH VIÊN", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13)
add_para("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13)
add_para("SỬ DỤNG MÔ HÌNH NGÔN NGỮ LỚN VÀ RAG", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=13)

for _ in range(2):
    doc.add_paragraph()

# Student info
info_lines = [
    "Giảng viên hướng dẫn : TS. ……………………",
    "Sinh viên thực hiện  : Trần Văn Trung",
    "Mã sinh viên         : ……………………",
    "Lớp                  : ……………………",
    "Khóa                 : ……………………",
]
for line in info_lines:
    add_para(line, font_size=13)

for _ in range(4):
    doc.add_paragraph()

add_para("Hà Nội – 2026", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=14)

add_page_break()

# ═══════════════════════════════════════════════════════════
# LỜI CẢM ƠN
# ═══════════════════════════════════════════════════════════

add_heading_styled("LỜI CẢM ƠN", level=1)

add_para(
    "Trong quá trình học tập và thực hiện báo cáo môn học \"Mô hình ngôn ngữ lớn\", "
    "em đã nhận được sự hướng dẫn, giúp đỡ tận tình từ các thầy cô giáo trong Khoa Công nghệ "
    "Thông tin – Trường Đại học Giao thông Vận tải. Em xin gửi lời cảm ơn chân thành đến "
    "các thầy cô đã trang bị cho em những kiến thức nền tảng về xử lý ngôn ngữ tự nhiên, "
    "học sâu và đặc biệt là về các mô hình ngôn ngữ lớn (LLM) và kỹ thuật truy vấn tạo sinh "
    "tăng cường (RAG), làm cơ sở để em có thể thực hiện báo cáo này."
)

add_para(
    "Em cũng xin cảm ơn gia đình và bạn bè đã luôn động viên, hỗ trợ em trong suốt "
    "thời gian học tập và nghiên cứu. Mặc dù đã cố gắng hết sức, nhưng do hạn chế về "
    "thời gian và kiến thức, báo cáo không tránh khỏi những thiếu sót. Em rất mong nhận "
    "được sự góp ý của thầy cô để báo cáo được hoàn thiện hơn."
)

add_para("Em xin trân trọng cảm ơn!", bold=True)

add_page_break()

print(f"Đã viết: Trang bìa + Lời cảm ơn")
print("Đang viết tiếp... (script quá dài, cần chạy tuần tự)")
