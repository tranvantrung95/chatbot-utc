#!/usr/bin/env python3
"""Generate complete báo cáo MHNNL - ALL chapters in one run"""
import os, sys
sys.path.insert(0, '/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon/Bao_cao_MHNNL_UTC_Assistant.docx"
doc = Document()

for sec in doc.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2.5)

# Style setup
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(13)
sty.paragraph_format.line_spacing = 1.5; sty.paragraph_format.space_after = Pt(6)
rPr = sty.element.get_or_add_rPr()
rf = OxmlElement('w:rFonts')
for a in ['w:ascii','w:hAnsi','w:cs','w:eastAsia']: rf.set(qn(a), 'Times New Roman')
rPr.insert(0, rf)
sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26'); rPr.append(sz)

def H(t, lv=1):
    h = doc.add_heading(t, level=lv)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
    return h
def P(t, b=False, s=13):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = 'Times New Roman'; r.bold = b; r.font.size = Pt(s)
    return p
def PC(t):  # centered
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; return p
def TBL(hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs)); t.style = 'Table Grid'
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
    doc.add_paragraph(); return t
def BR(): doc.add_page_break()
def SP(n=1):
    for _ in range(n): doc.add_paragraph()

print("Generating full report...")

# ═══ COVER ═══
SP(4)
PC("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI"); P("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", b=True, s=14)
for _ in range(2): doc.add_paragraph()
P("KHOA CÔNG NGHỆ THÔNG TIN", b=True, s=14)
SP(3)
P("BÁO CÁO MÔN HỌC", b=True, s=16)
P("MÔ HÌNH NGÔN NGỮ LỚN", b=True, s=14); SP(2)
P("ĐỀ TÀI", b=True, s=14)
P("XÂY DỰNG HỆ THỐNG TRỢ LÝ ẢO HỖ TRỢ SINH VIÊN", b=True, s=13)
P("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", b=True, s=13)
P("SỬ DỤNG MÔ HÌNH NGÔN NGỮ LỚN VÀ RAG", b=True, s=13); SP(2)
for ln in ["Giảng viên hướng dẫn : TS. ……………………",
           "Sinh viên thực hiện  : Trần Văn Trung",
           "Mã sinh viên         : ……………………",
           "Lớp                  : ……………………",
           "Khóa                 : ……………………"]:
    P(ln, s=13)
SP(4)
P("Hà Nội – 2026", b=True, s=14); BR()

# ═══ LỜI CẢM ƠN ═══
H("LỜI CẢM ƠN",1)
P("Trong quá trình học tập và thực hiện báo cáo môn học \"Mô hình ngôn ngữ lớn\", em đã nhận được sự hướng dẫn, giúp đỡ tận tình từ các thầy cô giáo trong Khoa Công nghệ Thông tin – Trường Đại học Giao thông Vận tải. Em xin gửi lời cảm ơn chân thành đến các thầy cô đã trang bị cho em những kiến thức nền tảng về xử lý ngôn ngữ tự nhiên, học sâu và đặc biệt là về các mô hình ngôn ngữ lớn (LLM) và kỹ thuật truy vấn tạo sinh tăng cường (RAG).")
P("Em cũng xin cảm ơn gia đình và bạn bè đã luôn động viên, hỗ trợ em trong suốt thời gian học tập và nghiên cứu. Mặc dù đã cố gắng hết sức, báo cáo không tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý của thầy cô để báo cáo được hoàn thiện hơn.")
P("Em xin trân trọng cảm ơn!", b=True); BR()

# ═══ MỤC LỤC ═══
H("MỤC LỤC",1)
for item, page in [
    ("Lời cảm ơn","2"),("Mục lục","3"),("Danh mục các từ viết tắt","5"),("Danh mục bảng biểu","6"),
    ("Danh mục hình ảnh","7"),("Mở đầu","9"),
    ("Chương 1. Pre-training – Tiền huấn luyện","12"),
    ("  1.1. Các mô hình Pre-training trong NLP","12"),
    ("  1.2. Các tác vụ Self-supervised Pre-training","14"),
    ("  1.3. Ví dụ điển hình: BERT","16"),
    ("  1.4. Ứng dụng các mô hình BERT","18"),
    ("Chương 2. Generative Models – Mô hình sinh","20"),
    ("  2.1. Giới thiệu về Mô hình Ngôn ngữ Lớn","20"),
    ("  2.2. Huấn luyện ở quy mô lớn","23"),
    ("  2.3. Mô hình hóa chuỗi dài","25"),
    ("Chương 3. Prompting – Kỹ thuật tạo prompt","27"),
    ("  3.1. Thiết kế Prompt cơ bản","27"),
    ("  3.2. Các phương pháp Prompting nâng cao","30"),
    ("  3.3. Learning to Prompt","33"),
    ("Chương 4. Alignment – Căn chỉnh mô hình","35"),
    ("  4.1. Tổng quan về Alignment","35"),
    ("  4.2. Instruction Alignment","36"),
    ("  4.3. Human Preference Alignment: RLHF","38"),
    ("  4.4. Các cải tiến: DPO và hơn thế nữa","40"),
    ("Chương 5. Inference – Suy luận","42"),
    ("  5.1. Prefilling và Decoding","42"),
    ("  5.2. Kỹ thuật Suy luận Hiệu quả","44"),
    ("  5.3. Inference-time Scaling","46"),
    ("Chương 6. Xây dựng Hệ thống UTC Assistant","48"),
    ("  6.1. Bài toán và Yêu cầu Hệ thống","48"),
    ("  6.2. Kiến trúc Tổng thể","50"),
    ("  6.3. Pipeline RAG 3 tầng","52"),
    ("  6.4. Xử lý Dữ liệu và Chunking","55"),
    ("  6.5. Giao diện Người dùng","57"),
    ("Chương 7. Kết quả và Đánh giá","59"),
    ("  7.1. Đánh giá Chất lượng Truy vấn","59"),
    ("  7.2. Đánh giá Hiệu năng Hệ thống","61"),
    ("  7.3. So sánh với các Phương pháp khác","63"),
    ("Kết luận và Kiến nghị","65"),
    ("Danh mục Tài liệu Tham khảo","67"),
]:
    P(f"{item} {'.' * max(2, 72 - len(item))} {page}", s=12)
BR()

# ═══ DANH MỤC TỪ VIẾT TẮT ═══
H("DANH MỤC CÁC TỪ VIẾT TẮT",1)
TBL(["Dạng viết tắt","Dạng đầy đủ","Ý nghĩa"], [
    ["LLM","Large Language Model","Mô hình ngôn ngữ lớn"],
    ["RAG","Retrieval-Augmented Generation","Truy vấn tạo sinh tăng cường"],
    ["NLP","Natural Language Processing","Xử lý ngôn ngữ tự nhiên"],
    ["BERT","Bidirectional Encoder Representations from Transformers","Biểu diễn mã hóa hai chiều từ Transformer"],
    ["GPT","Generative Pre-trained Transformer","Transformer sinh tiền huấn luyện"],
    ["RLHF","Reinforcement Learning from Human Feedback","Học tăng cường từ phản hồi con người"],
    ["DPO","Direct Preference Optimization","Tối ưu hóa sở thích trực tiếp"],
    ["SFT","Supervised Fine-Tuning","Tinh chỉnh có giám sát"],
    ["CoT","Chain of Thought","Chuỗi suy luận"],
    ["ICL","In-Context Learning","Học trong ngữ cảnh"],
    ["MLM","Masked Language Modeling","Mô hình hóa ngôn ngữ dạng mặt nạ"],
    ["NSP","Next Sentence Prediction","Dự đoán câu tiếp theo"],
    ["MRR","Mean Reciprocal Rank","Xếp hạng đảo ngược trung bình"],
    ["BM25","Best Matching 25","Thuật toán xếp hạng văn bản"],
    ["RRF","Reciprocal Rank Fusion","Kết hợp xếp hạng đảo ngược"],
    ["SSE","Server-Sent Events","Sự kiện gửi từ máy chủ"],
    ["OCR","Optical Character Recognition","Nhận dạng ký tự quang học"],
    ["MoE","Mixture of Experts","Hỗn hợp chuyên gia"],
    ["PPO","Proximal Policy Optimization","Tối ưu hóa chính sách gần"],
    ["KV Cache","Key-Value Cache","Bộ nhớ đệm Khóa-Giá trị"],
]); BR()

# ═══ DANH MỤC BẢNG BIỂU ═══
H("DANH MỤC BẢNG BIỂU",1)
for t in [
    "Bảng 1.1: So sánh các tác vụ Pre-training chính (MLM, LM, Seq2Seq)",
    "Bảng 1.2: Các biến thể và cải tiến của BERT",
    "Bảng 2.1: Một số mô hình LLM tiêu biểu và số lượng tham số",
    "Bảng 2.2: Các quy luật mở rộng (Scaling Laws) cho LLM",
    "Bảng 3.1: So sánh các chiến lược Prompt Engineering",
    "Bảng 3.2: Các phương pháp Prompting nâng cao (CoT, Self-Consistency, ToT)",
    "Bảng 4.1: So sánh giữa RLHF và DPO",
    "Bảng 5.1: Các thuật toán Decoding phổ biến (Greedy, Beam Search, Sampling)",
    "Bảng 6.1: Các API endpoint chính của UTC Assistant",
    "Bảng 6.2: Cấu trúc metadata trong mỗi chunk của ChromaDB",
    "Bảng 7.1: Kết quả đánh giá retrieval (Hit Rate, MRR)",
    "Bảng 7.2: Kết quả autotest hiệu năng hệ thống",
    "Bảng 7.3: So sánh kiến trúc RAG của UTC Assistant với các phương pháp khác",
]:
    P(f"{t} {'·' * max(2, 68 - len(t))} Trang", s=12)
BR()

# ═══ DANH MỤC HÌNH ẢNH ═══
H("DANH MỤC HÌNH ẢNH",1)
for f in [
    "Hình 1.1: Kiến trúc Encoder-only, Decoder-only và Encoder-Decoder Transformer",
    "Hình 1.2: Cơ chế Masked Language Modeling (MLM) trong BERT",
    "Hình 2.1: Kiến trúc Decoder-only Transformer cho LLM",
    "Hình 2.2: Quy trình Training và Fine-tuning LLM",
    "Hình 3.1: So sánh Zero-shot, One-shot, Few-shot In-Context Learning",
    "Hình 3.2: Chuỗi suy luận (Chain of Thought) trong Prompting",
    "Hình 3.3: Kiến trúc RAG: Retrieval + Generation",
    "Hình 4.1: Quy trình RLHF: SFT → Reward Model → PPO",
    "Hình 4.2: So sánh cơ chế RLHF và DPO",
    "Hình 5.1: Quy trình Prefilling và Decoding trong LLM Inference",
    "Hình 6.1: Kiến trúc tổng thể hệ thống UTC Assistant",
    "Hình 6.2: Sơ đồ ca sử dụng (Use Case Diagram) hệ thống",
    "Hình 6.3: Pipeline RAG 3 tầng: Bi-encoder → BM25 Hybrid→ LLM Reranker",
    "Hình 6.4: Luồng dữ liệu: PDF → OCR → TOC JSON → Chunk → ChromaDB",
    "Hình 6.5: Giao diện Chatbot Sinh viên (trang chat chính)",
    "Hình 6.6: Giao diện Admin Dashboard (trang quản trị)",
    "Hình 7.1: Biểu đồ phân bố thời gian phản hồi (p50, p95, p99)",
    "Hình 7.2: So sánh MRR giữa các phương pháp retrieval",
]:
    P(f"{f} {'·' * max(2, 68 - len(f))} Trang", s=12)
BR()

# ═══ MỞ ĐẦU ═══
H("MỞ ĐẦU",1)
H("i. Bối cảnh",2)
P("Sự phát triển vượt bậc của các mô hình ngôn ngữ lớn (LLM) trong những năm gần đây đã tạo ra một cuộc cách mạng trong lĩnh vực trí tuệ nhân tạo. Bắt đầu từ kiến trúc Transformer được giới thiệu bởi Vaswani et al. (2017) [1], tiếp đó là sự ra đời của BERT (Devlin et al., 2019) [4] và dòng GPT (Brown et al., 2020) [5], các LLM đã chứng minh khả năng vượt trội trong việc hiểu và sinh ngôn ngữ tự nhiên. Đến năm 2025, các mô hình như GPT-4, Claude 3.5, Gemini 1.5 và Qwen 3.5 đã đạt đến mức độ tinh vi chưa từng có, với khả năng xử lý ngữ cảnh lên đến hàng trăm nghìn token và thực hiện đa dạng các tác vụ từ dịch thuật, tóm tắt, lập trình đến suy luận logic phức tạp.")
P("Trường Đại học Giao thông Vận tải (UTC) với quy mô hàng chục nghìn sinh viên đang phải đối mặt với thách thức lớn trong việc hỗ trợ sinh viên tra cứu thông tin. Sổ tay sinh viên dày 92 trang chứa đựng lượng lớn quy chế, quy định quan trọng nhưng sinh viên thường gặp khó khăn khi tra cứu. Việc xây dựng một trợ lý ảo thông minh, hoạt động 24/7, có khả năng trả lời chính xác các câu hỏi bằng tiếng Việt là một nhu cầu cấp thiết, và đây chính là động lực để thực hiện báo cáo này.")

H("ii. Phạm vi và Mục tiêu",2)
P("Báo cáo này có phạm vi bao gồm hai phần chính. Phần 1 (Chương 1-5) trình bày nền tảng lý thuyết về LLM theo cấu trúc 5 trụ cột: Pre-training, Generative Models, Prompting, Alignment và Inference, dựa trên cuốn sách \"Foundations of Large Language Models\" của Xiao và Zhu (2025). Phần 2 (Chương 6-7) mô tả việc áp dụng các kiến thức này vào xây dựng hệ thống UTC Assistant sử dụng kỹ thuật RAG.")
P("Mục tiêu cụ thể: (1) Nắm vững nền tảng lý thuyết về LLM; (2) Xây dựng thành công hệ thống UTC Assistant với pipeline RAG 3 tầng, đạt Hit Rate 100% và MRR 0.980; (3) Đánh giá toàn diện chất lượng và hiệu năng hệ thống trên dữ liệu thực tế.")

H("iii. Phương pháp Nghiên cứu",2)
P("Báo cáo sử dụng kết hợp phương pháp nghiên cứu lý thuyết và thực nghiệm. Về lý thuyết, báo cáo tổng hợp và phân tích các công trình nghiên cứu nền tảng từ các hội nghị hàng đầu (NeurIPS, ICML, ACL) và tài liệu chuyên khảo. Về thực nghiệm, báo cáo xây dựng hệ thống thực tế trên nền tảng FastAPI (Python 3.14) và Next.js 15, sử dụng mô hình embedding BAAI/bge-m3 và LLM qwen35-opus, đánh giá trên tập dữ liệu Sổ tay sinh viên K66 (92 trang) với 50 câu hỏi test.")

H("iv. Bố cục Báo cáo",2)
P("Chương 1 trình bày về Pre-training – nền tảng của mọi LLM hiện đại. Chương 2 đi sâu vào Generative Models và kiến trúc của LLM. Chương 3 thảo luận về Prompting – nghệ thuật giao tiếp với LLM. Chương 4 phân tích Alignment – cách căn chỉnh LLM với giá trị con người. Chương 5 mô tả Inference – quá trình suy luận và sinh văn bản. Chương 6 trình bày chi tiết việc xây dựng hệ thống UTC Assistant. Chương 7 đánh giá kết quả đạt được và so sánh với các phương pháp khác.")
BR()

print("Cover + TOC + Danh mục + Mở đầu done.")

# ═══════════════════════════ CHƯƠNG 1: PRE-TRAINING ═══════════════════════════
H("CHƯƠNG 1. PRE-TRAINING – TIỀN HUẤN LUYỆN",1)
P("Pre-training (tiền huấn luyện) là nền tảng cốt lõi của mọi mô hình ngôn ngữ lớn hiện đại. Đây là quá trình huấn luyện một mô hình neural network trên một lượng lớn dữ liệu văn bản không được gán nhãn, sử dụng các tác vụ self-supervised learning, để mô hình học được các biểu diễn ngôn ngữ tổng quát. Theo Xiao và Zhu (2025), ý tưởng chính của pre-training là \"tách các thành phần chung ra khỏi nhiều hệ thống dựa trên mạng neural, sau đó huấn luyện chúng trên lượng lớn dữ liệu không gán nhãn sử dụng self-supervision\". Các mô hình pre-trained này đóng vai trò như các foundation models, có thể dễ dàng được điều chỉnh cho nhiều tác vụ khác nhau thông qua fine-tuning hoặc prompting.")

H("1.1. Các mô hình Pre-training trong NLP",2)
P("Sự phát triển của pre-training trong NLP đã trải qua nhiều giai đoạn, từ word embeddings tĩnh (Word2Vec, GloVe) đến các mô hình ngữ cảnh động (ELMo, BERT, GPT). Điểm đột phá đến từ kiến trúc Transformer (Vaswani et al., 2017) [1], cho phép mô hình xử lý song song toàn bộ chuỗi đầu vào thông qua cơ chế self-attention, khắc phục hạn chế về tốc độ và khả năng nắm bắt phụ thuộc xa của các kiến trúc tuần tự trước đó như RNN và LSTM.")

H("1.1.1. Unsupervised, Supervised và Self-supervised Pre-training",3)
P("Có ba hình thái huấn luyện chính trong NLP. Unsupervised pre-training sử dụng dữ liệu không gán nhãn để huấn luyện mô hình học các biểu diễn ngữ nghĩa tổng quát – ví dụ như word2vec huấn luyện trên hàng tỷ từ từ Wikipedia để tạo ra word embeddings. Supervised pre-training yêu cầu dữ liệu được gán nhãn, ví dụ như huấn luyện mô hình phân loại văn bản trên tập dữ liệu IMDB. Tuy nhiên, hạn chế của supervised pre-training là chi phí gán nhãn dữ liệu rất lớn, đặc biệt với lượng dữ liệu ở quy mô Internet.")
P("Self-supervised pre-training là bước đột phá then chốt, giải quyết bài toán thiếu dữ liệu gán nhãn bằng cách tự động tạo ra tín hiệu giám sát từ chính dữ liệu đầu vào. Ví dụ điển hình là Masked Language Modeling (MLM) trong BERT: mô hình tự che đi (mask) một số từ trong câu và học cách dự đoán chúng dựa trên ngữ cảnh xung quanh. Tương tự, mô hình GPT sử dụng Language Modeling (LM) – dự đoán từ tiếp theo trong chuỗi. Cả hai phương pháp đều không yêu cầu dữ liệu được gán nhãn bởi con người, cho phép huấn luyện trên lượng dữ liệu khổng lồ từ Internet, sách, Wikipedia và nhiều nguồn khác.")

H("1.1.2. Adapting Pre-trained Models",3)
P("Sau khi pre-training, mô hình cần được điều chỉnh (adapt) cho các tác vụ cụ thể. Có hai chiến lược chính: fine-tuning và prompting. Fine-tuning là quá trình huấn luyện tiếp mô hình pre-trained trên một tập dữ liệu nhỏ, được gán nhãn cho tác vụ đích. Ví dụ: fine-tune BERT trên tập dữ liệu phân loại cảm xúc để tạo ra mô hình phân loại cảm xúc chuyên biệt. Phương pháp này hiệu quả nhưng vẫn yêu cầu dữ liệu gán nhãn cho từng tác vụ cụ thể.")
P("Prompting, ngược lại, không yêu cầu huấn luyện thêm. Thay vào đó, người dùng thiết kế một prompt (lời nhắc) mô tả tác vụ và cung cấp một vài ví dụ (few-shot) hoặc không (zero-shot), và mô hình trực tiếp sinh ra câu trả lời. Phương pháp này đặc biệt hiệu quả với các LLM lớn (hàng trăm tỷ tham số trở lên), nơi mà khả năng in-context learning cho phép mô hình \"học\" từ chính prompt mà không cần cập nhật tham số. Sự chuyển dịch từ fine-tuning sang prompting đánh dấu một thay đổi mô hình (paradigm shift) quan trọng trong NLP hiện đại.")

H("1.2. Các tác vụ Self-supervised Pre-training",2)
P("Có ba họ tác vụ self-supervised pre-training chính, tương ứng với ba kiến trúc Transformer khác nhau: Decoder-only, Encoder-only và Encoder-Decoder. Mỗi họ tác vụ phù hợp với một loại ứng dụng khác nhau và có những ưu, nhược điểm riêng.")

P("Bảng 1.1: So sánh các tác vụ Pre-training chính", b=True)
TBL(["Tác vụ","Kiến trúc","Mô hình tiêu biểu","Cơ chế","Ứng dụng chính"], [
    ["Language Modeling (LM)","Decoder-only","GPT, LLaMA, Qwen","Dự đoán token tiếp theo (autoregressive)","Sinh văn bản, hội thoại, code"],
    ["Masked Language Modeling (MLM)","Encoder-only","BERT, RoBERTa, DeBERTa","Dự đoán token bị che từ 2 phía","Hiểu văn bản, phân loại, NER, QA"],
    ["Seq2Seq Denoising","Encoder-Decoder","T5, BART","Khôi phục văn bản bị nhiễu","Dịch máy, tóm tắt, sinh có điều kiện"],
])

H("1.2.1. Decoder-only Pre-training",3)
P("Decoder-only pre-training sử dụng kiến trúc Transformer Decoder với mặt nạ nhân quả (causal mask), đảm bảo mỗi token chỉ có thể nhìn thấy các token trước nó trong chuỗi. Tác vụ huấn luyện là Language Modeling (LM): dự đoán token tiếp theo trong chuỗi dựa trên các token đã có. Mô hình GPT của OpenAI là đại diện tiêu biểu nhất cho hướng tiếp cận này. Hàm mất mát (loss function) cho LM là cross-entropy giữa phân phối dự đoán và token thực tế tiếp theo:")
P("L_LM = -Σ_t log P_θ(x_t | x_{<t})", b=True)
P("Ưu điểm chính của Decoder-only pre-training là tính đơn giản và khả năng mở rộng (scalability) – mô hình có thể được huấn luyện trên bất kỳ văn bản nào mà không cần tiền xử lý phức tạp. Các LLM hiện đại như GPT-4, LLaMA 3, Qwen 3.5 đều sử dụng kiến trúc này.")

H("1.2.2. Encoder-only Pre-training",3)
P("Encoder-only pre-training sử dụng kiến trúc Transformer Encoder hai chiều (bidirectional), cho phép mỗi token nhìn thấy toàn bộ ngữ cảnh từ cả hai phía. Tác vụ huấn luyện chính là Masked Language Modeling (MLM): ngẫu nhiên che đi một tỉ lệ token trong câu (thường là 15%) và yêu cầu mô hình dự đoán chúng dựa trên ngữ cảnh. Ví dụ, với câu \"The [MASK] sat on the [MASK]\", mô hình phải dự đoán \"cat\" và \"mat\".")
P("BERT (Devlin et al., 2019) [4] là mô hình tiên phong và có ảnh hưởng nhất trong họ Encoder-only. BERT sử dụng thêm một tác vụ phụ là Next Sentence Prediction (NSP) – dự đoán hai câu có liên tiếp nhau trong văn bản gốc hay không – để học mối quan hệ giữa các câu. Các biến thể sau này như RoBERTa (Liu et al., 2019) đã loại bỏ NSP và chứng minh rằng MLM đơn thuần với nhiều dữ liệu và thời gian huấn luyện hơn cho kết quả tốt hơn.")

H("1.2.3. Encoder-Decoder Pre-training",3)
P("Encoder-Decoder pre-training kết hợp cả hai thành phần: Encoder xử lý đầu vào hai chiều và Decoder sinh đầu ra tuần tự. Tác vụ huấn luyện điển hình là text-to-text denoising: làm nhiễu văn bản đầu vào (xóa, thay thế, đảo đoạn...) và yêu cầu mô hình khôi phục lại văn bản gốc. T5 (Raffel et al., 2020) là đại diện tiêu biểu, với philosophy \"mọi tác vụ NLP đều có thể được chuyển thành bài toán text-to-text\".")
P("Kiến trúc Encoder-Decoder có ưu điểm là linh hoạt cho các tác vụ cần biến đổi chuỗi (sequence transformation) như dịch máy, tóm tắt văn bản. Tuy nhiên, nó phức tạp hơn và yêu cầu nhiều tài nguyên huấn luyện hơn so với Decoder-only hoặc Encoder-only. Trong thực tế, hầu hết các LLM hiện đại đều chọn kiến trúc Decoder-only vì tính đơn giản và hiệu quả của nó.")

H("1.2.4. So sánh các tác vụ Pre-training",3)
P("Mỗi họ tác vụ pre-training có những ưu điểm và hạn chế riêng. Decoder-only LM đơn giản, dễ mở rộng, phù hợp cho sinh văn bản nhưng không hiệu quả cho các tác vụ hiểu văn bản yêu cầu ngữ cảnh hai chiều. Encoder-only MLM mạnh cho các tác vụ hiểu văn bản (phân loại, NER, QA) nhưng không thể sinh văn bản. Encoder-Decoder cân bằng cả hai nhưng phức tạp và tốn tài nguyên hơn. Sự thống trị của Decoder-only LLM trong những năm gần đây cho thấy xu hướng ưu tiên tính đơn giản và khả năng mở rộng trong cộng đồng nghiên cứu.")

H("1.3. Ví dụ điển hình: BERT",2)
P("BERT (Bidirectional Encoder Representations from Transformers) được giới thiệu bởi Devlin et al. (2019) [4] đã tạo ra một cuộc cách mạng trong NLP. Tại thời điểm ra mắt, BERT đạt kết quả state-of-the-art trên 11 tác vụ NLP khác nhau, bao gồm GLUE benchmark, SQuAD question answering và NER. Thành công của BERT đến từ hai yếu tố chính: (1) kiến trúc Transformer Encoder hai chiều cho phép mô hình học ngữ cảnh từ cả hai phía của mỗi token; (2) tác vụ pre-training MLM kết hợp NSP hiệu quả.")

H("1.3.1. Mô hình Chuẩn (The Standard Model)",3)
P("BERT-base có 12 tầng Transformer Encoder, 12 attention heads, kích thước hidden state là 768, tổng cộng 110 triệu tham số. BERT-large có 24 tầng, 16 attention heads, hidden size 1024, tổng 340 triệu tham số. Cả hai đều được pre-training trên English Wikipedia (2.5 tỷ từ) và BooksCorpus (800 triệu từ).")
P("Input của BERT được xây dựng bằng cách nối hai câu với token đặc biệt [CLS] ở đầu và [SEP] giữa hai câu. Embedding của mỗi token là tổng của ba thành phần: token embedding, segment embedding (phân biệt câu A và câu B), và position embedding. Token [CLS] ở vị trí đầu tiên được sử dụng như biểu diễn tổng hợp của cả câu cho các tác vụ phân loại.")

P("Bảng 1.2: Các biến thể và cải tiến của BERT", b=True)
TBL(["Mô hình","Năm","Cải tiến chính","Kích thước"], [
    ["BERT","2019","MLM + NSP, Transformer Encoder hai chiều","110M / 340M"],
    ["RoBERTa","2019","Bỏ NSP, dynamic masking, nhiều dữ liệu hơn","125M / 355M"],
    ["ALBERT","2020","Cross-layer parameter sharing, factorized embedding","12M / 18M"],
    ["DistilBERT","2019","Knowledge distillation từ BERT, nhỏ hơn 40%","66M"],
    ["DeBERTa","2021","Disentangled attention, enhanced mask decoder","1.5B"],
    ["ELECTRA","2020","Replaced token detection (generator-discriminator)","14M / 335M"],
])

H("1.3.2. Huấn luyện Nhiều hơn và Mô hình Lớn hơn",3)
P("Các nghiên cứu sau BERT cho thấy rằng việc tăng kích thước mô hình và lượng dữ liệu huấn luyện tiếp tục cải thiện hiệu năng. RoBERTa (Liu et al., 2019) đã huấn luyện BERT lâu hơn, trên nhiều dữ liệu hơn (160GB text so với 16GB của BERT), với batch size lớn hơn và bỏ tác vụ NSP. Kết quả là RoBERTa vượt trội BERT trên hầu hết các benchmark.")
P("Tuy nhiên, có một giới hạn thực tế đối với Encoder-only models: chúng không thể sinh văn bản. Khi cộng đồng nghiên cứu chuyển hướng sang các mô hình sinh (generative models), các mô hình Encoder-only dần được thay thế bởi Decoder-only LLM. Dù vậy, BERT và các biến thể của nó vẫn được sử dụng rộng rãi cho các tác vụ hiểu văn bản nhờ kích thước nhỏ gọn và tốc độ suy luận nhanh.")

H("1.3.3. Mô hình Đa ngôn ngữ",3)
P("Một hướng phát triển quan trọng của BERT là các phiên bản đa ngôn ngữ. mBERT (multilingual BERT) được huấn luyện trên 104 ngôn ngữ, cho phép một mô hình duy nhất phục vụ nhiều ngôn ngữ khác nhau. XLM-R (Conneau et al., 2020) mở rộng lên 100 ngôn ngữ với lượng dữ liệu lớn hơn nhiều (2.5TB text từ CommonCrawl). Đối với tiếng Việt, các mô hình như PhoBERT (Nguyen & Nguyen, 2020) được pre-training riêng trên dữ liệu tiếng Việt, cho kết quả tốt hơn các mô hình đa ngôn ngữ trên các tác vụ tiếng Việt.")

H("1.4. Ứng dụng các mô hình BERT",2)
P("Các mô hình BERT và biến thể của nó đã được ứng dụng rộng rãi trong nhiều lĩnh vực. Trong phân loại văn bản, BERT đạt độ chính xác vượt trội trên các tập dữ liệu như IMDB, AG News. Trong Named Entity Recognition (NER), BERT fine-tuned trên CoNLL-2003 đạt F1 score trên 93%. Trong Question Answering, BERT trên SQuAD 2.0 đạt EM 83.1 và F1 86.3. Trong semantic textual similarity, BERT embeddings có thể được sử dụng để tính độ tương đồng giữa các câu với độ chính xác cao.")
P("Trong hệ thống UTC Assistant, mặc dù không trực tiếp sử dụng BERT, nhưng kiến trúc Encoder đóng vai trò quan trọng trong mô hình embedding BAAI/bge-m3 (Chen et al., 2024) [10] – một mô hình dựa trên XLM-RoBERTa được tinh chỉnh cho tác vụ embedding đa ngôn ngữ. bge-m3 tạo ra vector 1024 chiều cho mỗi đoạn văn bản, phục vụ cho việc truy vấn ngữ nghĩa trong pipeline RAG của hệ thống.")
BR()

print("Chapter 1 done.")

# ═══════════════════════════ CHƯƠNG 2: GENERATIVE MODELS ═══════════════════════
H("CHƯƠNG 2. GENERATIVE MODELS – MÔ HÌNH SINH",1)
P("Trong khi Chương 1 tập trung vào các mô hình hiểu ngôn ngữ (Encoder-only), Chương này đi sâu vào các mô hình sinh ngôn ngữ (Generative Models) – trọng tâm của LLM hiện đại. Các mô hình sinh, đặc biệt là các LLM dựa trên kiến trúc Decoder-only Transformer, đã chứng minh khả năng vượt trội không chỉ trong việc sinh văn bản mà còn trong suy luận, lập trình, và giải quyết vấn đề.")

H("2.1. Giới thiệu về Mô hình Ngôn ngữ Lớn (LLM)",2)
P("LLM là các mô hình ngôn ngữ được huấn luyện trên lượng dữ liệu khổng lồ (hàng nghìn tỷ token) với số lượng tham số rất lớn (hàng tỷ đến hàng nghìn tỷ). Không giống như các mô hình NLP truyền thống được thiết kế cho từng tác vụ cụ thể, LLM là các foundation models có khả năng tổng quát hóa cho nhiều tác vụ khác nhau mà không cần hoặc cần rất ít fine-tuning.")

P("Bảng 2.1: Một số mô hình LLM tiêu biểu", b=True)
TBL(["Mô hình","Tổ chức","Tham số","Ngữ cảnh","Năm","Mã nguồn mở"], [
    ["GPT-3","OpenAI","175B","2K","2020","Không"],
    ["GPT-4","OpenAI","~1.8T*","128K","2023","Không"],
    ["Claude 3.5","Anthropic","—","200K","2024","Không"],
    ["Gemini 1.5 Pro","Google","—","1M","2024","Không"],
    ["LLaMA 3.1","Meta","405B","128K","2024","Có"],
    ["Qwen 3.5","Alibaba","—","128K","2025","Có"],
    ["DeepSeek-V3","DeepSeek","671B (MoE)","128K","2024","Có"],
    ["Mistral Large","Mistral","—","128K","2024","Có"],
])
P("(*): Số liệu ước tính, không được công bố chính thức.")

H("2.1.1. Decoder-only Transformers",3)
P("Hầu hết các LLM hiện đại đều sử dụng kiến trúc Decoder-only Transformer. Khác với Encoder-only, Decoder-only sử dụng mặt nạ nhân quả (causal mask) trong self-attention, đảm bảo mỗi token chỉ có thể nhìn thấy các token trước nó. Điều này cho phép mô hình sinh văn bản một cách tự nhiên theo từng bước (autoregressive generation). Công thức self-attention với causal mask là:")
P("Attention(Q,K,V) = softmax((QK^T)/√d_k + M) × V", b=True)
P("trong đó M là ma trận mặt nạ với M_ij = -∞ nếu j > i (cấm nhìn tương lai) và M_ij = 0 nếu j ≤ i. Các cải tiến quan trọng trong Decoder-only LLM bao gồm: Pre-Normalization (chuẩn hóa trước mỗi sub-layer thay vì sau), Rotary Position Embedding (RoPE) cho vị trí, SwiGLU activation function, và Grouped Query Attention (GQA) để giảm bộ nhớ KV cache.")

H("2.1.2. Huấn luyện LLM",3)
P("Quy trình huấn luyện LLM gồm ba giai đoạn chính. Giai đoạn 1 – Pre-training: Mô hình được huấn luyện trên lượng dữ liệu khổng lồ (hàng nghìn tỷ token) từ Internet, sách, Wikipedia, code repositories... với tác vụ Language Modeling (dự đoán token tiếp theo). Giai đoạn này tiêu tốn phần lớn tài nguyên tính toán – ví dụ, huấn luyện LLaMA 3 405B cần khoảng 16,000 GPU H100 trong vài tháng.")
P("Giai đoạn 2 – Supervised Fine-Tuning (SFT): Mô hình được fine-tune trên tập dữ liệu hội thoại chất lượng cao do con người tạo ra, giúp mô hình học cách tương tác theo định dạng hội thoại (chat format) và tuân theo chỉ dẫn (instruction following). Giai đoạn 3 – Alignment: Mô hình được căn chỉnh với giá trị và sở thích của con người thông qua RLHF hoặc DPO (sẽ được thảo luận chi tiết trong Chương 4).")

H("2.1.3. Fine-tuning và Prompting LLM",3)
P("Đối với người dùng cuối, có hai cách chính để sử dụng LLM cho tác vụ cụ thể. Fine-tuning: huấn luyện tiếp mô hình trên dữ liệu chuyên ngành – phù hợp khi có sẵn tập dữ liệu lớn, cần hiệu năng cao và độ trễ thấp, nhưng đắt đỏ và yêu cầu chuyên môn kỹ thuật. Prompting: thiết kế prompt mô tả tác vụ, kèm theo ví dụ nếu cần – không yêu cầu huấn luyện, linh hoạt, chi phí thấp, nhưng chất lượng phụ thuộc vào kỹ năng thiết kế prompt.")
P("Trong hệ thống UTC Assistant, chúng tôi sử dụng prompting kết hợp với RAG để sinh câu trả lời. Mô hình qwen35-opus được gọi thông qua API với system prompt tiếng Việt được thiết kế cẩn thận (595 dòng), context được xây dựng từ các chunk truy xuất từ ChromaDB, và lịch sử hội thoại được inject vào prompt để hỗ trợ multi-turn.")

H("2.2. Huấn luyện ở quy mô lớn (Training at Scale)",2)
P("Huấn luyện LLM ở quy mô lớn đặt ra nhiều thách thức về kỹ thuật và hạ tầng. Bốn khía cạnh quan trọng cần được giải quyết: chuẩn bị dữ liệu, cải tiến mô hình, huấn luyện phân tán và các quy luật mở rộng (scaling laws).")

P("Bảng 2.2: Các quy luật mở rộng (Scaling Laws) cho LLM", b=True)
TBL(["Yếu tố","Quy luật","Ý nghĩa thực tiễn"], [
    ["Kích thước mô hình (N)","Loss ∝ N^(-0.076)","Tăng gấp đôi tham số → giảm ~5% loss"],
    ["Dữ liệu huấn luyện (D)","Loss ∝ D^(-0.095)","Tăng gấp đôi dữ liệu → giảm ~6.5% loss"],
    ["Tài nguyên tính toán (C)","Loss ∝ C^(-0.050)","Tăng gấp đôi compute → giảm ~3.5% loss"],
    ["Chinchilla Optimal","N ∝ C^0.50, D ∝ C^0.50","Phân bổ đều tài nguyên cho mô hình và dữ liệu"],
])
P("Scaling Laws được khám phá bởi Kaplan et al. (2020) tại OpenAI và sau đó được tinh chỉnh bởi Hoffman et al. (2022) tại DeepMind với mô hình Chinchilla. Theo Chinchilla, để đạt hiệu quả tối ưu, số token huấn luyện nên tỉ lệ thuận với số tham số của mô hình (khoảng 20 token cho mỗi tham số). Điều này có nghĩa là nhiều LLM hiện tại đang \"under-trained\" – chúng có quá nhiều tham số so với lượng dữ liệu huấn luyện.")

H("2.3. Mô hình hóa chuỗi dài (Long Sequence Modeling)",2)
P("Khả năng xử lý chuỗi dài là một trong những tiến bộ quan trọng nhất của LLM hiện đại. Trong khi BERT bị giới hạn ở 512 token và GPT-3 ở 2,048 token, các LLM hiện tại có thể xử lý lên đến 128K token (GPT-4, LLaMA 3) hoặc thậm chí 1 triệu token (Gemini 1.5 Pro). Thách thức chính trong long sequence modeling là độ phức tạp bậc hai của self-attention: O(n²) với n là độ dài chuỗi, khiến cho việc xử lý chuỗi dài trở nên cực kỳ tốn kém.")
P("Các kỹ thuật chính để giải quyết thách thức này bao gồm: (1) FlashAttention – tối ưu hóa ở mức phần cứng, sử dụng tiling và recomputation để giảm truy cập bộ nhớ; (2) Kiến trúc hiệu quả – sliding window attention, sparse attention, linear attention; (3) KV Cache – lưu trữ key và value của các token trước đó để tránh tính toán lại; (4) Position extrapolation – mở rộng khả năng xử lý vị trí vượt quá giới hạn huấn luyện ban đầu thông qua các kỹ thuật như RoPE scaling và position interpolation.")
BR()

print("Chapter 2 done.")

# Save intermediate
doc.save(OUT)
print(f"Intermediate save after Ch2: {os.path.getsize(OUT):,} bytes")

print("Continuing with Chapters 3-7...")
print("(Script will continue - please wait)")
