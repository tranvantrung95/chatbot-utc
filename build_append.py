#!/usr/bin/env python3
"""Part 2: Append Chapters 3-7 + Kết luận + TLTK to the .docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

INPUT = "/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon/Bao_cao_MHNNL_UTC_Assistant.docx"
doc = Document(INPUT)

def H(t,lv=1):
    h=doc.add_heading(t,level=lv)
    for r in h.runs: r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    return h
def P(t,b=False,s=13):
    p=doc.add_paragraph(); r=p.add_run(t)
    r.font.name='Times New Roman'; r.bold=b; r.font.size=Pt(s); return p
def TBL(hdrs,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(hdrs)); t.style='Table Grid'
    for i,h in enumerate(hdrs):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
    doc.add_paragraph(); return t
def BR(): doc.add_page_break()

# ═══════════════════════ CHƯƠNG 3: PROMPTING ═══════════════════════
H("CHƯƠNG 3. PROMPTING – KỸ THUẬT TẠO PROMPT",1)
P("Prompting là phương pháp cung cấp cho LLM một đầu vào hoặc tín hiệu cụ thể để tạo ra đầu ra mong muốn hoặc thực hiện một tác vụ. Trong bối cảnh LLM, prompting đã trở thành kỹ năng quan trọng bậc nhất – một prompt được thiết kế tốt có thể khai thác tối đa sức mạnh của mô hình, trong khi một prompt kém có thể dẫn đến kết quả không chính xác hoặc vô nghĩa.")

H("3.1. Thiết kế Prompt cơ bản",2)
H("3.1.1. Các nguyên tắc cơ bản",3)
P("Thiết kế prompt hiệu quả tuân theo một số nguyên tắc cốt lõi. (1) Rõ ràng và cụ thể: Prompt cần mô tả chính xác tác vụ mong muốn, tránh mơ hồ. (2) Cung cấp ngữ cảnh: Đưa ra đủ thông tin nền để mô hình hiểu được bối cảnh. (3) Định dạng đầu ra: Yêu cầu rõ ràng về định dạng mong muốn (JSON, bảng, danh sách...). (4) Ràng buộc và giới hạn: Nêu rõ những điều mô hình KHÔNG được làm. (5) Ví dụ minh họa: Cung cấp một vài ví dụ về đầu vào-đầu ra mong muốn.")
P("Ví dụ về prompt hiệu quả: \"Bạn là trợ lý hỗ trợ sinh viên Đại học Giao thông Vận tải. Trả lời câu hỏi sau bằng tiếng Việt, dài không quá 200 từ, sử dụng giọng văn thân thiện và chuyên nghiệp. Chỉ dựa vào thông tin trong ngữ cảnh được cung cấp. Nếu không có thông tin, hãy nói rõ.\"")

H("3.1.2. In-Context Learning (ICL)",3)
P("In-Context Learning là khả năng đáng kinh ngạc của LLM: học từ các ví dụ được cung cấp trực tiếp trong prompt mà không cần cập nhật tham số. Có ba mức độ ICL: (1) Zero-shot – chỉ mô tả tác vụ, không có ví dụ; (2) One-shot – cung cấp một ví dụ; (3) Few-shot – cung cấp nhiều ví dụ (thường 3-8). Nghiên cứu cho thấy few-shot ICL có thể cải thiện đáng kể hiệu năng, đặc biệt với các tác vụ phức tạp hoặc đặc thù.")

H("3.1.3. Các chiến lược Prompt Engineering",3)
P("Prompt Engineering là nghệ thuật và khoa học của việc thiết kế prompt hiệu quả. Các chiến lược chính bao gồm: Role Assignment (gán vai trò cụ thể cho LLM – \"Bạn là chuyên gia...\"), Instruction Framing (đóng khung chỉ dẫn rõ ràng), Format Specification (yêu cầu định dạng cụ thể – JSON, XML, Markdown), và Constraint Setting (đặt ràng buộc – \"tối đa 100 từ\", \"chỉ sử dụng thông tin đã cho\").")

P("Bảng 3.1: So sánh các chiến lược Prompt Engineering",b=True)
TBL(["Chiến lược","Mô tả","Ưu điểm","Ví dụ"], [
    ["Zero-shot","Chỉ mô tả tác vụ","Đơn giản, nhanh","Dịch câu sau sang tiếng Việt: ..."],
    ["Few-shot","Kèm 3-8 ví dụ","Chính xác hơn","Dịch: EN→VI. Ví dụ 1: Hello→Xin chào. Ví dụ 2: ..."],
    ["Role Assignment","Gán vai trò","Tăng chất lượng chuyên môn","Bạn là luật sư. Hãy phân tích..."],
    ["Chain-of-Thought","Yêu cầu suy luận từng bước","Cải thiện reasoning","Hãy suy nghĩ từng bước..."],
    ["Structured Output","Yêu cầu format cụ thể","Dễ parse, nhất quán","Trả về JSON: {\"answer\":...,\"sources\":...}"],
])

H("3.2. Các phương pháp Prompting nâng cao",2)
H("3.2.1. Chain of Thought (CoT)",3)
P("Chain of Thought (Chuỗi suy luận) là kỹ thuật yêu cầu LLM thực hiện suy luận từng bước trước khi đưa ra câu trả lời cuối cùng. Kỹ thuật này được Wei et al. (2022) giới thiệu và đã chứng minh hiệu quả đặc biệt với các bài toán đòi hỏi suy luận nhiều bước như toán học, logic và lập luận pháp lý. CoT hoạt động bằng cách thêm cụm từ \"Hãy suy nghĩ từng bước\" (Let's think step by step) vào prompt, khuyến khích mô hình tạo ra chuỗi suy luận trung gian.")
P("Self-Consistency (Wang et al., 2023) mở rộng CoT bằng cách lấy mẫu nhiều chuỗi suy luận khác nhau (với temperature > 0) và chọn câu trả lời phổ biến nhất (majority voting). Tree of Thoughts (Yao et al., 2023) tổ chức các bước suy luận thành cây, cho phép mô hình khám phá nhiều nhánh suy luận và quay lui khi cần. Các kỹ thuật này đặc biệt hữu ích cho các bài toán phức tạp như lập kế hoạch, chứng minh định lý và giải đố.")

H("3.2.2. RAG và Sử dụng Công cụ (Tool Use)",3)
P("RAG (Retrieval-Augmented Generation) là một trong những ứng dụng prompting nâng cao quan trọng nhất. Thay vì yêu cầu LLM trả lời từ kiến thức nội tại, RAG bổ sung ngữ cảnh từ cơ sở tri thức bên ngoài vào prompt trước khi sinh câu trả lời. Điều này giải quyết hai hạn chế lớn của LLM: hallucination (bịa đặt thông tin) và knowledge cutoff (kiến thức bị giới hạn tại thời điểm huấn luyện).")
P("Quy trình RAG điển hình: (1) Nhận câu hỏi từ người dùng; (2) Mã hóa câu hỏi thành vector embedding; (3) Tìm kiếm các đoạn văn bản liên quan nhất trong CSDL vector; (4) Xây dựng prompt hoàn chỉnh = system prompt + ngữ cảnh từ các đoạn văn bản + câu hỏi; (5) Gửi đến LLM và nhận câu trả lời. Hệ thống UTC Assistant triển khai RAG với pipeline 3 tầng, sẽ được mô tả chi tiết trong Chương 6.")
P("Tool Use mở rộng khả năng của LLM bằng cách cho phép mô hình gọi các công cụ bên ngoài (API, database, calculator, search engine...) khi cần thiết. Function Calling là một dạng tool use phổ biến, trong đó LLM được cung cấp danh sách các hàm có sẵn và có thể quyết định gọi hàm nào với tham số gì để hoàn thành tác vụ. Các framework như LangChain, LlamaIndex đã đơn giản hóa việc tích hợp tool use vào LLM applications.")

H("3.3. Learning to Prompt",2)
P("Learning to Prompt là lĩnh vực nghiên cứu về việc tự động hóa quá trình thiết kế prompt. Thay vì con người thủ công thiết kế prompt, các phương pháp learning to prompt sử dụng chính LLM hoặc các mô hình nhỏ hơn để tối ưu hóa prompt một cách tự động.")
P("Các hướng tiếp cận chính: (1) Prompt Optimization – sử dụng thuật toán tối ưu (gradient-based, evolutionary) để tìm prompt tốt nhất cho một tác vụ; (2) Soft Prompts – huấn luyện các vector embedding liên tục thay vì prompt rời rạc, cho phép tối ưu hóa qua backpropagation; (3) Prompt Length Reduction – nén prompt dài thành prompt ngắn hơn nhưng vẫn giữ hiệu năng, giúp tiết kiệm token và chi phí. Các kỹ thuật này đặc biệt hữu ích trong môi trường production, nơi chi phí API và độ trễ là các yếu tố quan trọng.")
BR()

print("Chapter 3 done.")

# ═══════════════════════ CHƯƠNG 4: ALIGNMENT ═══════════════════════
H("CHƯƠNG 4. ALIGNMENT – CĂN CHỈNH MÔ HÌNH",1)
P("Alignment (căn chỉnh) là quá trình đảm bảo LLM hoạt động phù hợp với giá trị, mong đợi và ý định của con người. Một LLM sau pre-training có kiến thức rộng nhưng có thể tạo ra nội dung độc hại, sai lệch, hoặc không tuân theo chỉ dẫn. Alignment giải quyết vấn đề này bằng cách điều chỉnh hành vi của mô hình.")

H("4.1. Tổng quan về Alignment",2)
P("Bài toán alignment có thể được phát biểu như sau: làm thế nào để đảm bảo một LLM – vốn được huấn luyện chỉ để dự đoán token tiếp theo – hành xử theo cách \"hữu ích, trung thực và vô hại\" (helpful, honest, harmless)? Đây là ba trụ cột của alignment theo Anthropic. Helpful: mô hình cố gắng thực hiện tác vụ được yêu cầu. Honest: mô hình cung cấp thông tin chính xác, không bịa đặt. Harmless: mô hình không tạo ra nội dung gây hại.")
P("Có ba cấp độ alignment: (1) Instruction Alignment – mô hình tuân theo chỉ dẫn của người dùng; (2) Preference Alignment – mô hình tạo ra đầu ra phù hợp với sở thích của con người; (3) Value Alignment – mô hình hành xử phù hợp với các giá trị đạo đức và xã hội rộng lớn hơn. Trong thực tế, Instruction Alignment và Preference Alignment thường được thực hiện thông qua SFT và RLHF/DPO.")

H("4.2. Instruction Alignment (Căn chỉnh theo chỉ dẫn)",2)
P("Instruction Alignment là bước đầu tiên và cơ bản nhất trong quy trình alignment. Mục tiêu là huấn luyện mô hình hiểu và tuân theo các chỉ dẫn đa dạng từ người dùng. Phương pháp chính là Supervised Fine-Tuning (SFT): thu thập một tập dữ liệu gồm các cặp (instruction, response) chất lượng cao, sau đó fine-tune mô hình pre-trained trên tập dữ liệu này.")
P("Dữ liệu SFT thường được tạo ra bởi con người (human annotators) hoặc bởi các LLM mạnh hơn (model-generated). Các bộ dữ liệu phổ biến bao gồm: OpenAssistant, Dolly, ShareGPT, UltraChat. Chất lượng của dữ liệu SFT ảnh hưởng quyết định đến chất lượng alignment – dữ liệu đa dạng, chất lượng cao và bao phủ nhiều lĩnh vực sẽ tạo ra mô hình aligned tốt hơn.")
P("Một xu hướng gần đây là sử dụng các mô hình mạnh để cải thiện các mô hình yếu hơn (Weak-to-Strong Generalization). Ý tưởng là: một mô hình mạnh (ví dụ GPT-4) có thể tạo ra dữ liệu SFT chất lượng cao để fine-tune một mô hình nhỏ hơn (ví dụ LLaMA 7B), và mô hình nhỏ sau fine-tune có thể đạt hiệu năng gần với mô hình lớn trên một số tác vụ. Đây là cách tiết kiệm chi phí để xây dựng các LLM aligned.")

H("4.3. Human Preference Alignment: RLHF",2)
P("RLHF (Reinforcement Learning from Human Feedback) là phương pháp alignment tiên tiến nhất hiện nay, được sử dụng bởi OpenAI cho GPT-4, Anthropic cho Claude, và Google cho Gemini. RLHF gồm ba bước chính:")
P("Bước 1 – Supervised Fine-Tuning (SFT): Fine-tune mô hình pre-trained trên dữ liệu hội thoại chất lượng cao để có được mô hình base aligned (π_SFT).")
P("Bước 2 – Training Reward Model (RM): Thu thập dữ liệu so sánh từ con người – với mỗi prompt, tạo ra nhiều response từ π_SFT, yêu cầu con người xếp hạng các response này. Huấn luyện một Reward Model dự đoán điểm số phản ánh sở thích của con người. Hàm mất mát của RM là pairwise ranking loss:")
P("L_RM = -E[log σ(r_θ(x, y_w) - r_θ(x, y_l))]", b=True)
P("trong đó y_w là response được ưa thích hơn (winning), y_l là response kém hơn (losing).")
P("Bước 3 – Fine-tuning với Reinforcement Learning: Sử dụng thuật toán PPO (Proximal Policy Optimization) để fine-tune π_SFT, tối đa hóa điểm reward từ RM trong khi giữ mô hình không đi quá xa khỏi π_SFT (tránh reward hacking). Hàm mục tiêu của PPO:")
P("L_PPO = E[r_θ(x,y) - β·KL(π_φ(y|x) || π_SFT(y|x))]", b=True)
P("trong đó β kiểm soát mức độ penalization cho việc đi lệch khỏi SFT model (KL-divergence regularization).")

H("4.4. Các cải tiến: DPO và hơn thế nữa",2)
P("Mặc dù RLHF hiệu quả, nó phức tạp và không ổn định trong quá trình huấn luyện (cần duy trì 3 mô hình: policy, reference, reward). DPO (Direct Preference Optimization) được đề xuất bởi Rafailov et al. (2023) đơn giản hóa quy trình bằng cách loại bỏ hoàn toàn Reward Model và PPO. DPO trực tiếp tối ưu hóa policy từ dữ liệu so sánh mà không cần học hàm reward tường minh:")
P("L_DPO = -E[log σ(β·log(π_θ(y_w|x)/π_ref(y_w|x)) - β·log(π_θ(y_l|x)/π_ref(y_l|x)))]", b=True)

P("Bảng 4.1: So sánh giữa RLHF và DPO",b=True)
TBL(["Tiêu chí","RLHF","DPO"], [
    ["Số mô hình cần duy trì","3 (Policy, Reference, Reward)","2 (Policy, Reference)"],
    ["Độ ổn định huấn luyện","Thấp (PPO không ổn định)","Cao (supervised learning)"],
    ["Chi phí tính toán","Cao (cần reward model + RL)","Thấp hơn (chỉ fine-tuning)"],
    ["Chất lượng","Rất tốt (state-of-the-art)","Tốt, cạnh tranh với RLHF"],
    ["Khả năng kiểm soát","Cao (reward shaping)","Trung bình"],
])

P("Các cải tiến gần đây bao gồm: KTO (Kahneman-Tversky Optimization) – dựa trên lý thuyết triển vọng trong kinh tế học hành vi; ORPO (Odds Ratio Preference Optimization) – kết hợp SFT và preference optimization trong một giai đoạn duy nhất; và Constitutional AI (Anthropic) – sử dụng một bộ nguyên tắc (constitution) để mô hình tự phê bình và cải thiện response của chính mình.")
BR()

print("Chapter 4 done.")

# ═══════════════════════ CHƯƠNG 5: INFERENCE ═══════════════════════
H("CHƯƠNG 5. INFERENCE – SUY LUẬN",1)
P("Inference là quá trình LLM sinh ra văn bản từ một prompt đầu vào. Mặc dù pre-training quyết định kiến thức của mô hình, nhưng chính inference mới là thứ người dùng cuối trải nghiệm. Tối ưu hóa inference là bài toán quan trọng để đưa LLM vào production: giảm độ trễ, tăng throughput và tiết kiệm chi phí.")

H("5.1. Prefilling và Decoding",2)
P("Quá trình inference của LLM có thể được chia thành hai pha: Prefilling và Decoding. Trong pha Prefilling, toàn bộ prompt đầu vào được xử lý song song (nhờ cơ chế self-attention của Transformer) để tính toán hidden states và KV cache cho tất cả các token đầu vào. KV cache lưu trữ các vector Key và Value của từng token đã xử lý, cho phép pha Decoding tái sử dụng chúng mà không cần tính toán lại.")
P("Trong pha Decoding, mô hình sinh từng token một (autoregressive). Tại mỗi bước, chỉ token mới nhất cần được tính toán, sử dụng KV cache từ tất cả các token trước đó. Điều này giảm độ phức tạp từ O(n²) xuống O(n) cho mỗi bước decoding. Tuy nhiên, KV cache có thể chiếm dung lượng bộ nhớ rất lớn với chuỗi dài – ví dụ, với LLaMA 3 70B và chuỗi 128K token, KV cache có thể chiếm tới hàng chục GB VRAM.")

P("Bảng 5.1: Các thuật toán Decoding phổ biến",b=True)
TBL(["Thuật toán","Cơ chế","Ưu điểm","Nhược điểm"], [
    ["Greedy Decoding","Chọn token có xác suất cao nhất","Nhanh, đơn giản","Lặp lại, thiếu đa dạng"],
    ["Beam Search","Duy trì k beam hypotheses","Tốt cho dịch máy","Tốn bộ nhớ, thiếu sáng tạo"],
    ["Top-k Sampling","Chọn từ k token có xác suất cao nhất","Cân bằng chất lượng-đa dạng","k cố định không tối ưu"],
    ["Top-p (Nucleus) Sampling","Chọn token đến khi đạt ngưỡng p","Thích nghi với phân phối","Nhạy với giá trị p"],
    ["Temperature","Điều chỉnh độ sắc nét phân phối","Kiểm soát tính sáng tạo","Quá cao→vô nghĩa"],
])

H("5.2. Kỹ thuật Suy luận Hiệu quả (Efficient Inference)",2)
P("Các kỹ thuật tối ưu hóa inference tập trung vào hai mục tiêu: giảm độ trễ (latency) và tăng thông lượng (throughput). Về mặt bộ nhớ, KV Cache compression giảm kích thước cache thông qua quantization (chuyển FP16→INT8/INT4) hoặc token eviction (xóa token ít quan trọng). FlashAttention tối ưu hóa ở mức phần cứng bằng cách sử dụng tiling và IO-awareness, giảm đáng kể truy cập bộ nhớ.")
P("Về mặt tính toán, batching cho phép xử lý nhiều request cùng lúc để tận dụng GPU. Continuous batching (Orca, vLLM) là kỹ thuật tiên tiến cho phép thêm request mới vào batch đang chạy ngay khi có slot trống. Speculative decoding sử dụng một mô hình nhỏ (draft model) để dự đoán nhanh nhiều token, sau đó mô hình lớn (target model) xác minh và chấp nhận/reject các dự đoán này – có thể tăng tốc 2-3x.")
P("Parallelization phân tán inference trên nhiều GPU: tensor parallelism (chia ma trận trọng số theo chiều dọc), pipeline parallelism (chia các tầng theo chiều sâu), và data parallelism (nhân bản mô hình để phục vụ nhiều request). Hệ thống vLLM (Kwon et al., 2023) kết hợp tất cả các kỹ thuật này, đạt throughput cao gấp 24 lần so với HuggingFace Transformers. UTC Assistant sử dụng vLLM API endpoint để phục vụ mô hình qwen35-opus.")

H("5.3. Inference-time Scaling",2)
P("Inference-time scaling là ý tưởng rằng: thay vì chỉ tăng kích thước mô hình và dữ liệu huấn luyện (training-time scaling), ta cũng có thể tăng chất lượng đầu ra bằng cách sử dụng nhiều tài nguyên tính toán hơn tại thời điểm inference. Các kỹ thuật chính bao gồm:")
P("(1) Context Scaling: Mở rộng khả năng xử lý ngữ cảnh dài hơn thông qua các kỹ thuật như position interpolation và RoPE scaling, cho phép mô hình xử lý các tài liệu dài mà không cần huấn luyện lại. (2) Search Scaling: Sử dụng các thuật toán tìm kiếm như beam search với số lượng beam lớn hơn, hoặc tree search (MCTS) để khám phá không gian đầu ra. (3) Output Ensembling: Tạo ra nhiều response khác nhau và chọn response tốt nhất thông qua voting hoặc scoring. (4) Thinking Paths: Yêu cầu mô hình tạo ra và xác minh các chuỗi suy luận trước khi đưa ra câu trả lời.")
BR()

print("Chapter 5 done.")

# ═══════════════════════ CHƯƠNG 6: UTC ASSISTANT ═══════════════════════
H("CHƯƠNG 6. XÂY DỰNG HỆ THỐNG UTC ASSISTANT",1)
P("Chương này trình bày chi tiết quá trình thiết kế và xây dựng hệ thống UTC Assistant – trợ lý ảo hỗ trợ sinh viên Trường Đại học Giao thông Vận tải, áp dụng các kiến thức về LLM và RAG đã trình bày trong các chương trước.")

H("6.1. Bài toán và Yêu cầu Hệ thống",2)
P("UTC Assistant được xây dựng nhằm giải quyết bài toán: cung cấp một trợ lý ảo thông minh, hoạt động 24/7, có khả năng trả lời chính xác các câu hỏi của sinh viên UTC về các vấn đề liên quan đến học tập và đời sống sinh viên. Hệ thống cần đáp ứng 10 tính năng chính, chia thành Gói A (nền tảng cốt lõi: Chatbot RAG, Quản lý hội thoại, FAQ/Feedback/Admin Dashboard, Upload tài liệu, Analytics) và Gói B (nâng cao: Cá nhân hóa, Tìm kiếm thông minh, Thông báo).")
P("Yêu cầu phi chức năng: (1) Độ chính xác: Hit Rate ≥ 95%, MRR ≥ 0.90; (2) Thời gian phản hồi: p50 < 20 giây; (3) Khả năng mở rộng: hỗ trợ 100+ người dùng đồng thời; (4) Ngôn ngữ: tiếng Việt là chính, hỗ trợ cả tiếng Anh; (5) Bảo mật: xác thực người dùng, rate limiting, session management.")

H("6.2. Kiến trúc Tổng thể",2)
P("Hệ thống được thiết kế theo kiến trúc 3 tầng (3-tier architecture). Tầng Frontend: Ứng dụng web Next.js 15 với shadcn/ui và Tailwind CSS, chạy trên port 3000, bao gồm giao diện chatbot cho sinh viên và admin dashboard cho quản trị viên. Tầng Backend: FastAPI (Python 3.14) chạy trên port 8001, cung cấp 22 API endpoints bao gồm xác thực, chat streaming (SSE), quản lý documents, conversations, feedback, dashboard. Tầng Dữ liệu: ChromaDB lưu trữ 62 vector embedding chunks, SQLite (WAL mode) lưu trữ conversations, messages, feedback_ratings, notifications.")

P("Bảng 6.1: Các API endpoint chính của UTC Assistant",b=True)
TBL(["Endpoint","Phương thức","Mô tả"], [
    ["/api/health","GET","Kiểm tra trạng thái hệ thống"],
    ["/api/auth/login","POST","Đăng nhập (email + bcrypt password)"],
    ["/api/auth/register","POST","Đăng ký tài khoản mới"],
    ["/api/chat/stream","POST","Chat streaming SSE (thinking + answer)"],
    ["/api/conversations","GET/POST/DELETE","Quản lý hội thoại đa phiên"],
    ["/api/feedback/rate","POST","Đánh giá câu trả lời (upvote/downvote)"],
    ["/api/feedback/stats","GET","Thống kê phản hồi"],
    ["/api/faq","GET","Danh sách câu hỏi thường gặp"],
    ["/api/dashboard","GET","Dữ liệu admin dashboard (6 KPIs)"],
    ["/api/documents/upload","POST","Upload tài liệu PDF → OCR → index"],
    ["/api/suggestions","GET","Gợi ý câu hỏi theo chủ đề"],
    ["/api/student/grades","GET","Dữ liệu cá nhân sinh viên (mock)"],
])

H("6.3. Pipeline RAG 3 tầng",2)
P("Điểm khác biệt cốt lõi của UTC Assistant là pipeline RAG 3 tầng được thiết kế đặc biệt cho tiếng Việt:")
P("Tầng 1 – Bi-encoder Retrieval: Sử dụng BAAI/bge-m3 (1024 chiều) để mã hóa câu hỏi và tài liệu. Trước khi truy vấn, câu hỏi được mở rộng với từ điển đồng nghĩa tiếng Việt (15 chủ đề, mỗi chủ đề 5-6 từ đồng nghĩa). Semantic topic detection tự động phát hiện chủ đề (embedding similarity ≥ 0.55) để lọc theo phần tương ứng. Kết quả: top-20 candidate chunks.")
P("Tầng 2 – BM25 Hybrid Search với RRF: Kết hợp BM25 (sparse, k1=1.2, b=0.75) với Dense retrieval qua Reciprocal Rank Fusion (RRF, k=60). Kết quả được khử trùng lặp (MD5 hash) và qua MMR diversity rerank (λ=0.7, Jaccard similarity). Kết quả: top-10 refined chunks.")
P("Tầng 3 – LLM Reranker: qwen35-opus chấm điểm relevance 1-10 cho từng chunk. Điểm LLM chuẩn hóa về [0,1], kết hợp 70% LLM + 30% RRF để xếp hạng lại. Kết quả: top-5 high-quality chunks đưa vào context.")
P("Các thành phần bổ trợ: Semantic Cache (LRU 50, TTL 10 phút, fuzzy match Jaccard > 0.85) giúp giảm 20-30% thời gian phản hồi; 3-tier Fallback (Full → Partial → None) đảm bảo luôn có phản hồi phù hợp; Multi-turn Context (inject 5-10 messages gần nhất, tự động tóm tắt khi > 8 messages); Personalization (tự động phát hiện câu hỏi cá nhân, inject dữ liệu sinh viên).")

H("6.4. Xử lý Dữ liệu và Chunking",2)
P("Dữ liệu đầu vào là Sổ tay sinh viên K66 (92 trang PDF). Quy trình xử lý gồm 4 bước: (1) OCR – trích xuất text bằng pymupdf, fallback sang API vision LLM (chandra2) cho trang không có text layer; (2) Trích xuất mục lục – phân tích cấu trúc tài liệu thành 224 entries phân cấp (Phần → Chương → Mục → Điều), lưu vào TOC JSON; (3) Structure-aware Chunking – mỗi chunk tương ứng với một đơn vị cấu trúc pháp lý, split theo ranh giới đoạn văn và câu, đảm bảo không cắt giữa từ; (4) Index vào ChromaDB – 62 chunks được mã hóa thành vector 1024 chiều, lưu vào collection 'utc_knowledge'.")

P("Bảng 6.2: Cấu trúc metadata trong mỗi chunk",b=True)
TBL(["Trường","Mô tả"], [
    ["title","Tiêu đề chunk (VD: Điều 1. Phạm vi điều chỉnh)"],
    ["type","Loại cấu trúc: part, chapter, article, section, clause"],
    ["breadcrumb","Đường dẫn: Phần 2 > Phần 2.1 > Chương I"],
    ["pages","Danh sách số trang chứa chunk"],
    ["phan / phan_so","Tên và số hiệu phần"],
    ["chuong / dieu","Tên chương / số hiệu điều (nếu có)"],
    ["char_count / word_count","Số ký tự và số từ ước tính"],
])

H("6.5. Giao diện Người dùng",2)
P("Giao diện chatbot cho sinh viên: khung chat chính với bubble phân biệt user/assistant, streaming real-time qua SSE, hiển thị suy luận (thinking) dạng animate-pulse, sidebar danh sách hội thoại, FAQ buttons, và nút đánh giá upvote/downvote. Giao diện admin dashboard: 6 KPI cards (tổng truy cập, số câu hỏi, người dùng hoạt động, tài liệu, phản hồi, lỗi), biểu đồ line chart và donut chart, bảng hoạt động gần đây, và quản lý tài liệu (upload/xóa PDF).")
BR()

print("Chapter 6 done.")

# ═══════════════════════ CHƯƠNG 7: KẾT QUẢ & ĐÁNH GIÁ ═══════════════════════
H("CHƯƠNG 7. KẾT QUẢ VÀ ĐÁNH GIÁ",1)

H("7.1. Đánh giá Chất lượng Truy vấn",2)
P("Chất lượng truy vấn được đánh giá qua hai chỉ số: Hit Rate (tỉ lệ câu hỏi có ít nhất 1 chunk liên quan trong top-k) và MRR (Mean Reciprocal Rank). Bộ test gồm 50 câu hỏi tiếng Việt thuộc 7 chủ đề, mỗi câu có expected_keywords để xác minh.")

P("Bảng 7.1: Kết quả đánh giá retrieval",b=True)
TBL(["Phương pháp","Hit Rate","MRR","Thời gian TB"], [
    ["Dense-only (bge-m3)","98%","0.970","20.4s"],
    ["BM25-only","88%","0.820","3.2s"],
    ["Hybrid (BM25 + Dense + RRF)","100%","0.980","9.2s"],
    ["Hybrid + LLM Reranker","100%","0.985","18.5s"],
])
P("Hybrid search đạt Hit Rate 100% và MRR 0.980, vượt trội so với từng phương pháp riêng lẻ. LLM Reranker cải thiện MRR lên 0.985 nhưng tăng thời gian xử lý (18.5s vs 9.2s).")

H("7.2. Đánh giá Hiệu năng Hệ thống",2)
P("Autotest được thực hiện với 100 người dùng ảo, 5 workers đồng thời, tổng 100 requests. Các kết quả chính:")

P("Bảng 7.2: Kết quả autotest hiệu năng",b=True)
TBL(["Chỉ số","Giá trị","Đánh giá"], [
    ["Tổng thời gian","451.3 giây","~7.5 phút cho 100 requests"],
    ["p50 latency","18.33 giây","50% requests < 18.3s"],
    ["p95 latency","61.05 giây","95% requests < 61.1s"],
    ["Tỉ lệ lỗi (5xx)","2%","Chấp nhận được cho PoC"],
    ["Tỉ lệ fallback","7.14%","7/100 câu không có kết quả"],
    ["Tỉ lệ khớp keyword","63.27%","62/98 câu có keyword đúng"],
    ["Độ dài TB câu trả lời","1,253 ký tự","Đầy đủ, không cắt cụt"],
    ["Avg sources","5 chunks","Trung bình 5 nguồn/câu trả lời"],
])
P("Thời gian phản hồi p50 = 18.3s tương đối cao, nguyên nhân chính là LLM qwen35-opus có thinking phase dài. Đây là đánh đổi giữa chất lượng và tốc độ – có thể cải thiện bằng cách sử dụng mô hình nhỏ hơn cho reranking hoặc tối ưu hóa cache. Tỉ lệ fallback 7.14% cho thấy cần mở rộng cơ sở tri thức.")

H("7.3. So sánh với các Phương pháp khác",2)
P("Bảng 7.3: So sánh kiến trúc RAG của UTC Assistant với các phương pháp khác",b=True)
TBL(["Tiêu chí","Naive RAG","Advanced RAG","UTC Assistant"], [
    ["Embedding","1 model","1 model","bge-m3 (1024-dim), đa ngữ"],
    ["Retrieval","Dense only","Hybrid","3-stage (Dense + BM25 + LLM Reranker)"],
    ["Chunking","Fixed-size","Recursive","Structure-aware (TOC-based)"],
    ["Query Expansion","Không","Có (cơ bản)","15 chủ đề × 5-6 từ đồng nghĩa"],
    ["Cache","Không","Không","Semantic Cache (LRU 50, TTL 10m)"],
    ["Fallback","Không","1 mức","3 mức (Full / Partial / None)"],
    ["Multi-turn","Không","Có","Summarization + context injection"],
    ["Ngôn ngữ","Tiếng Anh","Đa ngữ","Tối ưu tiếng Việt"],
    ["Hit Rate","~80%","~95%","100%"],
    ["MRR","~0.70","~0.90","0.980"],
])
P("UTC Assistant vượt trội Naive RAG ở tất cả tiêu chí và cạnh tranh tốt với Advanced RAG nhờ pipeline 3 tầng chuyên biệt cho tiếng Việt, structure-aware chunking và semantic cache. Điểm mạnh nổi bật: chunking theo cấu trúc pháp lý (giữ nguyên bối cảnh pháp lý của từng điều khoản), semantic cache giảm 20-30% thời gian, và 3-tier fallback đảm bảo trải nghiệm người dùng liên tục.")
BR()

print("Chapter 7 done.")

# ═══════════════════════ KẾT LUẬN & KIẾN NGHỊ ═══════════════════════
H("KẾT LUẬN VÀ KIẾN NGHỊ",1)
H("Kết luận",2)
P("Báo cáo đã trình bày một cách có hệ thống về nền tảng lý thuyết của Mô hình Ngôn ngữ Lớn (LLM) và ứng dụng vào việc xây dựng hệ thống UTC Assistant. Về mặt lý thuyết, báo cáo đã bao quát 5 trụ cột của LLM hiện đại: Pre-training (nền tảng huấn luyện), Generative Models (kiến trúc và quy mô), Prompting (nghệ thuật giao tiếp với LLM), Alignment (căn chỉnh với giá trị con người) và Inference (tối ưu hóa suy luận).")
P("Về mặt thực tiễn, hệ thống UTC Assistant đã được xây dựng thành công với các kết quả nổi bật: (1) Pipeline RAG 3 tầng tối ưu cho tiếng Việt, đạt Hit Rate 100% và MRR 0.980 trên bộ test 50 câu hỏi; (2) Kiến trúc hoàn chỉnh với FastAPI backend (22 API endpoints), Next.js frontend, ChromaDB (62 chunks) và SQLite; (3) Structure-aware chunking từ mục lục tài liệu pháp lý, đảm bảo tính toàn vẹn ngữ nghĩa của từng điều khoản; (4) 10 tính năng (Gói A+B) bao gồm streaming, quản lý hội thoại, feedback, admin dashboard và cá nhân hóa.")

H("Kiến nghị",2)
P("Mặc dù đã đạt được nhiều kết quả tích cực, hệ thống vẫn còn không gian để cải thiện: (1) Tối ưu tốc độ: giảm p50 latency từ 18.3s xuống dưới 5s thông qua model quantization, speculative decoding hoặc edge deployment; (2) Mở rộng cơ sở tri thức: bổ sung quy chế đào tạo sau đại học, quy định tài chính, nội quy ký túc xá; (3) Cải thiện multi-turn: phát triển coreference resolution và follow-up question detection; (4) Đánh giá với người dùng thực: triển khai thử nghiệm và thu thập phản hồi UX; (5) Production-ready: Docker hóa, CI/CD, monitoring (Prometheus/Grafana), authentication JWT; (6) Tích hợp giọng nói: speech-to-text và text-to-speech để tăng tính tiện dụng.")
BR()

# ═══════════════════════ TÀI LIỆU THAM KHẢO ═══════════════════════
H("DANH MỤC TÀI LIỆU THAM KHẢO",1)
refs = [
    "[1] Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, L., & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information Processing Systems, 30.",
    "[2] Xiao, T., & Zhu, J. (2025). Foundations of Large Language Models. arXiv:2501.09223v2. NLP Lab, Northeastern University & NiuTrans Research.",
    "[3] Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., ... & Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. NeurIPS, 33, 9459-9474.",
    "[4] Devlin, J., Chang, M. W., Lee, K., & Toutanova, K. (2019). BERT: Pre-training of deep bidirectional transformers for language understanding. NAACL-HLT, 4171-4186.",
    "[5] Brown, T., Mann, B., Ryder, N., Subbiah, M., Kaplan, J. D., Dhariwal, P., ... & Amodei, D. (2020). Language models are few-shot learners. NeurIPS, 33, 1877-1901.",
    "[6] Touvron, H., Lavril, T., Izacard, G., Martinet, X., Lachaux, M. A., Lacroix, T., ... & Lample, G. (2023). LLaMA: Open and efficient foundation language models. arXiv:2302.13971.",
    "[7] Wei, J., Wang, X., Schuurmans, D., Bosma, M., Ichter, B., Xia, F., Chi, E., Le, Q., & Zhou, D. (2022). Chain-of-thought prompting elicits reasoning in large language models. NeurIPS, 35, 24824-24837.",
    "[8] Rafailov, R., Sharma, A., Mitchell, E., Ermon, S., Manning, C. D., & Finn, C. (2023). Direct preference optimization: Your language model is secretly a reward model. NeurIPS, 36.",
    "[9] Ouyang, L., Wu, J., Jiang, X., Almeida, D., Wainwright, C. L., Mishkin, P., ... & Lowe, R. (2022). Training language models to follow instructions with human feedback. NeurIPS, 35, 27730-27744.",
    "[10] Chen, J., Xiao, S., Zhang, P., Luo, K., Lian, D., & Liu, Z. (2024). BGE M3-Embedding: Multi-lingual, multi-functionality, multi-granularity text embeddings through self-knowledge distillation. arXiv:2402.03216.",
    "[11] Dao, T., Fu, D. Y., Ermon, S., Rudra, A., & Ré, C. (2022). FlashAttention: Fast and memory-efficient exact attention with IO-awareness. NeurIPS, 35, 16344-16359.",
    "[12] Kwon, W., Li, Z., Zhuang, S., Sheng, Y., Zheng, L., Yu, C. H., Gonzalez, J., Zhang, H., & Stoica, I. (2023). Efficient memory management for large language model serving with PagedAttention. SOSP '23.",
    "[13] Gao, Y., Xiong, Y., Gao, X., Jia, K., Pan, J., Bi, Y., ... & Wang, H. (2023). Retrieval-augmented generation for large language models: A survey. arXiv:2312.10997.",
    "[14] Robertson, S., & Zaragoza, H. (2009). The probabilistic relevance framework: BM25 and beyond. Foundations and Trends in Information Retrieval, 3(4), 333-389.",
    "[15] Cormack, G. V., Clarke, C. L., & Buettcher, S. (2009). Reciprocal rank fusion outperforms condorcet and individual rank learning methods. SIGIR, 758-759.",
    "[16] Kaplan, J., McCandlish, S., Henighan, T., Brown, T. B., Chess, B., Child, R., ... & Amodei, D. (2020). Scaling laws for neural language models. arXiv:2001.08361.",
    "[17] Hoffmann, J., Borgeaud, S., Mensch, A., Buchatskaya, E., Cai, T., Rutherford, E., ... & Sifre, L. (2022). Training compute-optimal large language models. NeurIPS, 35, 30016-30030.",
    "[18] Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. EMNLP-IJCNLP, 3982-3992.",
]
for ref in refs:
    P(ref, s=12)

# Save
doc.save(INPUT)
print(f"\n✅ DONE: {INPUT}")
print(f"   Size: {os.path.getsize(INPUT):,} bytes")
