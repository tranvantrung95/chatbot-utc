#!/usr/bin/env python3
"""FULL rebuild: Bao_cao_MHNNL theo cấu trúc Foundation of LLM + UTC Assistant"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "/Users/trantrung/CAOHOC_CNTT/mo_hinh_ngon_ngu_lon/Bao_cao_MHNNL_UTC_Assistant.docx"
doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3); s.right_margin = Cm(2.5)

# ── Styles ──
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(13)
sty.paragraph_format.line_spacing = 1.5; sty.paragraph_format.space_after = Pt(6)
rPr = sty.element.get_or_add_rPr()
rf = OxmlElement('w:rFonts')
for a in ['w:ascii','w:hAnsi','w:cs','w:eastAsia']: rf.set(qn(a), 'Times New Roman')
rPr.insert(0, rf)
sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26'); rPr.append(sz)

def H(t, lv=1):
    h = doc.add_heading(t, level=lv)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
    return h

def P(t, b=False, al=None, s=13):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = 'Times New Roman'; r.bold = b; r.font.size = Pt(s)
    if al is not None: p.alignment = al
    return p

def TBL(hdrs, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs)); t.style = 'Table Grid'
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
    doc.add_paragraph(); return t

def BR(): doc.add_page_break()
def SP(n=1):
    for _ in range(n): doc.add_paragraph()

# ═══════════════════════════════════ TRANG BÌA ═══════════════════════════════════
SP(4)
P("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=14)
P("KHOA CÔNG NGHỆ THÔNG TIN", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=14); SP(3)
P("BÁO CÁO MÔN HỌC", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=16)
P("MÔ HÌNH NGÔN NGỮ LỚN", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=14); SP(2)
P("ĐỀ TÀI", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=14)
P("XÂY DỰNG HỆ THỐNG TRỢ LÝ ẢO HỖ TRỢ SINH VIÊN", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=13)
P("TRƯỜNG ĐẠI HỌC GIAO THÔNG VẬN TẢI", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=13)
P("SỬ DỤNG MÔ HÌNH NGÔN NGỮ LỚN VÀ RAG", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=13); SP(2)
for ln in ["Giảng viên hướng dẫn : TS. ……………………",
           "Sinh viên thực hiện  : Trần Văn Trung",
           "Mã sinh viên         : ……………………",
           "Lớp                  : ……………………",
           "Khóa                 : ……………………"]:
    P(ln, s=13)
SP(4)
P("Hà Nội – 2026", b=True, al=WD_ALIGN_PARAGRAPH.CENTER, s=14); BR()

# ═══════════════════════════════════ LỜI CẢM ƠN ═══════════════════════════════════
H("LỜI CẢM ƠN",1)
P("Trong quá trình học tập và thực hiện báo cáo môn học \"Mô hình ngôn ngữ lớn\", em đã nhận được sự hướng dẫn, giúp đỡ tận tình từ các thầy cô giáo trong Khoa Công nghệ Thông tin – Trường Đại học Giao thông Vận tải. Em xin gửi lời cảm ơn chân thành đến các thầy cô đã trang bị cho em những kiến thức nền tảng về xử lý ngôn ngữ tự nhiên, học sâu và đặc biệt là về các mô hình ngôn ngữ lớn (LLM) và kỹ thuật truy vấn tạo sinh tăng cường (RAG), làm cơ sở để em có thể thực hiện báo cáo này.")
P("Em cũng xin cảm ơn gia đình và bạn bè đã luôn động viên, hỗ trợ em trong suốt thời gian học tập và nghiên cứu. Mặc dù đã cố gắng hết sức, nhưng do hạn chế về thời gian và kiến thức, báo cáo không tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý của thầy cô để báo cáo được hoàn thiện hơn.")
P("Em xin trân trọng cảm ơn!", b=True); BR()

# ═══════════════════════════════════ MỤC LỤC ═══════════════════════════════════
H("MỤC LỤC",1)
toc_data = [
    ("Lời cảm ơn","2"),("Mục lục","3"),("Danh mục các từ viết tắt","5"),("Danh mục bảng biểu","6"),
    ("Danh mục hình ảnh","7"),("Mở đầu","9"),
    ("Chương 1. Pre-training – Tiền huấn luyện","12"),
    ("  1.1. Các mô hình Pre-training trong NLP","12"),
    ("  1.2. Các tác vụ Self-supervised Pre-training","15"),
    ("  1.3. Ví dụ điển hình: BERT","18"),
    ("  1.4. Ứng dụng các mô hình BERT","21"),
    ("Chương 2. Generative Models – Mô hình sinh","23"),
    ("  2.1. Giới thiệu về Mô hình Ngôn ngữ Lớn (LLM)","23"),
    ("  2.2. Huấn luyện ở quy mô lớn","27"),
    ("  2.3. Mô hình hóa chuỗi dài","30"),
    ("Chương 3. Prompting – Kỹ thuật tạo prompt","32"),
    ("  3.1. Thiết kế Prompt cơ bản","32"),
    ("  3.2. Các phương pháp Prompting nâng cao","36"),
    ("  3.3. Learning to Prompt","40"),
    ("Chương 4. Alignment – Căn chỉnh mô hình","42"),
    ("  4.1. Tổng quan về Alignment","42"),
    ("  4.2. Instruction Alignment","43"),
    ("  4.3. Human Preference Alignment: RLHF","45"),
    ("  4.4. Các cải tiến trong Human Preference Alignment","47"),
    ("Chương 5. Inference – Suy luận","49"),
    ("  5.1. Prefilling và Decoding","49"),
    ("  5.2. Kỹ thuật Suy luận Hiệu quả","51"),
    ("  5.3. Inference-time Scaling","53"),
    ("Chương 6. Xây dựng Hệ thống UTC Assistant","55"),
    ("  6.1. Bài toán và Yêu cầu Hệ thống","55"),
    ("  6.2. Kiến trúc Tổng thể","57"),
    ("  6.3. Pipeline RAG 3 tầng","59"),
    ("  6.4. Xử lý Dữ liệu và Chunking","63"),
    ("  6.5. Giao diện Người dùng","65"),
    ("Chương 7. Kết quả và Đánh giá","67"),
    ("  7.1. Đánh giá Chất lượng Truy vấn","67"),
    ("  7.2. Đánh giá Hiệu năng Hệ thống","69"),
    ("  7.3. So sánh với các Phương pháp khác","71"),
    ("Kết luận và Kiến nghị","73"),
    ("Danh mục Tài liệu Tham khảo","75"),
]
for item, page in toc_data:
    P(f"{item} {'.' * max(2, 70 - len(item))} {page}", s=12)
BR()

# ═══════════════════════════════════ DANH MỤC TỪ VIẾT TẮT ═══════════════════════
H("DANH MỤC CÁC TỪ VIẾT TẮT",1)
TBL(["Dạng viết tắt","Dạng đầy đủ","Ý nghĩa"], [
    ["LLM","Large Language Model","Mô hình ngôn ngữ lớn"],
    ["RAG","Retrieval-Augmented Generation","Truy vấn tạo sinh tăng cường"],
    ["NLP","Natural Language Processing","Xử lý ngôn ngữ tự nhiên"],
    ["BERT","Bidirectional Encoder Representations from Transformers","Biểu diễn mã hóa hai chiều từ Transformer"],
    ["GPT","Generative Pre-trained Transformer","Transformer sinh tiền huấn luyện"],
    ["RLHF","Reinforcement Learning from Human Feedback","Học tăng cường từ phản hồi con người"],
    ["DPO","Direct Preference Optimization","Tối ưu hóa sở thích trực tiếp"],
    ["SFT","Supervised Fine-Tuning","Tinh chỉnh có giám sát"],
    ["CoT","Chain of Thought","Chuỗi suy luận"],
    ["ICL","In-Context Learning","Học trong ngữ cảnh"],
    ["MLM","Masked Language Modeling","Mô hình hóa ngôn ngữ dạng mặt nạ"],
    ["NSP","Next Sentence Prediction","Dự đoán câu tiếp theo"],
    ["MRR","Mean Reciprocal Rank","Xếp hạng đảo ngược trung bình"],
    ["BM25","Best Matching 25","Thuật toán xếp hạng văn bản"],
    ["RRF","Reciprocal Rank Fusion","Kết hợp xếp hạng đảo ngược"],
    ["SSE","Server-Sent Events","Sự kiện gửi từ máy chủ"],
    ["OCR","Optical Character Recognition","Nhận dạng ký tự quang học"],
    ["API","Application Programming Interface","Giao diện lập trình ứng dụng"],
    ["MoE","Mixture of Experts","Hỗn hợp chuyên gia"],
    ["KV Cache","Key-Value Cache","Bộ nhớ đệm Khóa-Giá trị"],
]); BR()

# ═══════════════════════════════════ DANH MỤC BẢNG BIỂU ═════════════════════════
H("DANH MỤC BẢNG BIỂU",1)
for t in ["Bảng 1.1: So sánh các tác vụ Pre-training chính",
          "Bảng 1.2: Các biến thể và cải tiến của BERT",
          "Bảng 2.1: Một số mô hình LLM tiêu biểu",
          "Bảng 2.2: Các quy luật mở rộng (Scaling Laws) cho LLM",
          "Bảng 3.1: So sánh các chiến lược Prompt Engineering",
          "Bảng 3.2: Các phương pháp Prompting nâng cao",
          "Bảng 4.1: So sánh RLHF và DPO",
          "Bảng 5.1: Các thuật toán Decoding phổ biến",
          "Bảng 6.1: Các API endpoint chính của UTC Assistant",
          "Bảng 6.2: Cấu trúc metadata trong mỗi chunk",
          "Bảng 7.1: Kết quả đánh giá retrieval (Hit Rate, MRR)",
          "Bảng 7.2: Kết quả autotest hiệu năng",
          "Bảng 7.3: So sánh kiến trúc RAG của UTC Assistant với các phương pháp khác"]:
    P(f"{t} {'·' * max(2, 65 - len(t))} Trang", s=12)
BR()

# ═══════════════════════════════════ DANH MỤC HÌNH ẢNH ═════════════════════════
H("DANH MỤC HÌNH ẢNH",1)
for f in ["Hình 1.1: Kiến trúc Encoder-only, Decoder-only và Encoder-Decoder Transformer",
          "Hình 1.2: Cơ chế Masked Language Modeling (MLM) trong BERT",
          "Hình 2.1: Kiến trúc Decoder-only Transformer cho LLM",
          "Hình 2.2: Quy trình Training và Fine-tuning LLM",
          "Hình 3.1: So sánh Zero-shot, One-shot, Few-shot In-Context Learning",
          "Hình 3.2: Chuỗi suy luận (Chain of Thought) trong Prompting",
          "Hình 3.3: Kiến trúc RAG: Retrieval + Generation",
          "Hình 4.1: Quy trình RLHF: SFT → Reward Model → PPO",
          "Hình 4.2: So sánh RLHF và DPO",
          "Hình 5.1: Quy trình Prefilling và Decoding trong LLM Inference",
          "Hình 6.1: Kiến trúc tổng thể hệ thống UTC Assistant",
          "Hình 6.2: Sơ đồ ca sử dụng hệ thống",
          "Hình 6.3: Pipeline RAG 3 tầng: Bi-encoder → BM25 Hybrid → LLM Reranker",
          "Hình 6.4: Luồng dữ liệu: PDF → OCR → TOC JSON → Chunk → ChromaDB",
          "Hình 6.5: Giao diện Chatbot Sinh viên",
          "Hình 6.6: Giao diện Admin Dashboard",
          "Hình 7.1: Biểu đồ phân bố thời gian phản hồi (p50, p95, p99)",
          "Hình 7.2: So sánh MRR giữa các phương pháp retrieval"]:
    P(f"{f} {'·' * max(2, 65 - len(f))} Trang", s=12)
BR()

print("Part 1 done: Cover → Danh mục. Saving intermediate...")
doc.save(OUTPUT)
print(f"Intermediate save: {os.path.getsize(OUTPUT):,} bytes")
